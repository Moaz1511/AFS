import re
import tkinter as tk
from tkinter import filedialog, messagebox
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches
import pandas as pd
import json

# --- Constants and Configuration ---
FONT_NAME = "Tiro Bangla"
FONT_SIZE = 11
SOLUTION_HEADER_COLOR = RGBColor(0x08, 0x85, 0x65)
QUESTION_LABELS = ['ক', 'খ', 'গ', 'ঘ']

# --- Core Reusable Functions ---

def patch_omml_font_size(omml_xml, size_pt=FONT_SIZE):
    try:
        size_val = str(int(size_pt * 2))
        tree = etree.fromstring(omml_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for sz_tag in tree.xpath('.//w:sz|.//w:szCs', namespaces=ns):
            sz_tag.attrib[qn('w:val')] = size_val
        return etree.tostring(tree, encoding='unicode')
    except etree.XMLSyntaxError:
        return omml_xml

def set_bangla_font(run, font_size=FONT_SIZE, bold=False, color=None, underline=False):
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.underline = underline
    if color:
        run.font.color.rgb = color

def split_text_and_omml(para):
    para_xml = etree.tostring(para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    parts = []
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            run_text = "".join(node.xpath('.//w:t/text()', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
            if run_text:
                parts.append(("text", run_text))
        elif node.tag.endswith(('oMath', 'oMathPara')):
            omml_xml = etree.tostring(node, encoding='unicode')
            parts.append(("omml", omml_xml))
    return parts

def render_parts_to_para(para, parts_list, bold=False, color=None, underline=False):
    for ctype, cvalue in parts_list:
        if ctype == "text":
            run = para.add_run(cvalue)
            set_bangla_font(run, bold=bold, color=color, underline=underline)
        elif ctype == "omml":
            omml_patched = patch_omml_font_size(cvalue)
            omml_run = parse_xml(
                f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml_patched}</w:r>'
            )
            para._p.append(omml_run)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)

def get_para_full_text_with_omml(para):
    full_str = ""
    for ctype, cvalue in split_text_and_omml(para):
        full_str += cvalue if ctype == "text" else omml_to_latex_text(cvalue)
    return full_str.strip()

# --- Equation Parsing (OMML to Text) ---
MML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def to_unicode_math(text):
    def _math_italic(c):
        if 'a' <= c <= 'z': return chr(ord('\U0001D44E') + ord(c) - ord('a'))
        if 'A' <= c <= 'Z': return chr(ord('\U0001D434') + ord(c) - ord('A'))
        return c
    return ''.join(_math_italic(c) for c in text)

def _parse_omml_node_recursive(node):
    if node is None: return ""
    tag = etree.QName(node.tag).localname
    if tag == 't': return to_unicode_math(node.text or '')
    elif tag == 'f':
        num = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}num'))
        den = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}den'))
        return f"({num})/({den})"
    elif tag == 'sSup':
        base = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}e'))
        sup = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}sup'))
        return f"{base}^{{{sup}}}"
    elif tag == 'sSub':
        base = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}e'))
        sub = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}sub'))
        return f"{base}_{{{sub}}}"
    elif tag == 'rad':
        base = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}e'))
        return f"sqrt({base})"
    else:
        parts = [_parse_omml_node_recursive(child) for child in node]
        non_empty_parts = [p.strip() for p in parts if p and not p.isspace()]
        return " ".join(non_empty_parts)

def omml_to_latex_text(omml_xml):
    try:
        tree = etree.fromstring(omml_xml)
        math_node = tree.find(f'.//{{{MML_NS}}}oMath') or tree.find(f'.//{{{MML_NS}}}oMathPara')
        if math_node is None and etree.QName(tree.tag).localname in ('oMath', 'oMathPara'):
            math_node = tree
        return _parse_omml_node_recursive(math_node).strip() if math_node is not None else ""
    except Exception:
        return ""

# --- CQ Parsing Logic (REVISED) ---

def parse_stem_header(text):
    m = re.match(r'^\s*প্রশ্ন\s+([০-৯0-9]+)\.', text.strip())
    return m.group(1) if m else None

def parse_question_part(text):
    m = re.match(r'^\s*([কখগঘ])\.', text.strip())
    return m.group(1) if m else None

def parse_answer_part(text):
    m = re.match(r'^\s*উত্তর\s*[\(]?\s*([কখগঘ])\s*[\.\)]?', text.strip(), re.IGNORECASE)
    return m.group(1) if m else None

def extract_cqs(paragraphs, is_math=False):
    """Extracts all structured CQ data, including questions and answers."""
    cqs = []
    current_cq = None
    state = "find_stem"  # states: find_stem, in_question_block, in_solution_block
    current_solution_label = None
    q_labels = QUESTION_LABELS[:-1] if is_math else QUESTION_LABELS

    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue

        stem_serial = parse_stem_header(text)
        question_label = parse_question_part(text)
        answer_label = parse_answer_part(text)

        # Start of a new CQ
        if stem_serial:
            if current_cq:
                cqs.append(current_cq)
            current_cq = {'serial': stem_serial, 'stem_meta': [para], 'questions_meta': {}, 'solutions_meta': {}}
            state = "in_question_block"
            continue

        if not current_cq:
            continue

        if state == "in_question_block":
            if question_label and question_label in q_labels:
                current_cq['questions_meta'][question_label] = [para]
            elif answer_label and answer_label in q_labels:
                state = "in_solution_block"
                current_solution_label = answer_label
                current_cq['solutions_meta'][answer_label] = [para]
            else: # Continuation of the stem
                current_cq['stem_meta'].append(para)

        elif state == "in_solution_block":
            if answer_label and answer_label in q_labels:
                current_solution_label = answer_label
                current_cq['solutions_meta'][current_solution_label] = [para]
            elif current_solution_label: # Continuation of the current answer
                current_cq['solutions_meta'][current_solution_label].append(para)

    if current_cq: # Append the last CQ
        cqs.append(current_cq)

    return cqs

# --- Export and Formatting ---

def cqs_to_rows(cqs, is_math=False):
    """Converts a list of CQ dictionaries to rows for DataFrame export."""
    rows = []
    q_labels = QUESTION_LABELS[:-1] if is_math else QUESTION_LABELS

    for cq in cqs:
        row = {'Serial': cq['serial']}
        stem_text = " ".join(get_para_full_text_with_omml(p) for p in cq.get('stem_meta', []))
        row['Stem'] = re.sub(r'^\s*প্রশ্ন\s+[০-৯0-9]+\.\s*', '', stem_text).strip()

        for label in q_labels:
            q_paras = cq.get('questions_meta', {}).get(label, [])
            s_paras = cq.get('solutions_meta', {}).get(label, [])

            q_text = " ".join(get_para_full_text_with_omml(p) for p in q_paras)
            s_text = " ".join(get_para_full_text_with_omml(p) for p in s_paras)
            
            row[f'Question_{label}'] = re.sub(fr'^\s*{label}\.\s*', '', q_text).strip()
            row[f'Solution_{label}'] = re.sub(fr'^\s*উত্তর\s*[\(]?\s*{label}\s*[\.\)]?\s*', '', s_text, flags=re.IGNORECASE).strip()
            
        rows.append(row)
    return rows

def strip_prefix_from_parts(parts, prefix_regex):
    """
    Strips a prefix defined by a regex from a list of text/omml parts,
    even if the prefix is split across multiple text runs.
    """
    if not parts:
        return []

    # Reconstruct the initial text from the parts to match the regex against
    initial_text = ""
    for p_type, p_value in parts:
        if p_type == "text":
            initial_text += p_value
        else: # Stop if we hit a non-text element like an equation
            break
    
    match = re.match(prefix_regex, initial_text, re.IGNORECASE)
    if not match:
        return parts # No prefix found, return original parts

    chars_to_skip = len(match.group(0))
    new_parts = []
    skipped = False

    for p_type, p_value in parts:
        if skipped or p_type == "omml":
            new_parts.append((p_type, p_value))
            continue

        if p_type == "text":
            if len(p_value) <= chars_to_skip:
                chars_to_skip -= len(p_value)
                if chars_to_skip == 0:
                    skipped = True
            else:
                new_value = p_value[chars_to_skip:]
                new_parts.append((p_type, new_value))
                chars_to_skip = 0
                skipped = True
    
    return new_parts

def format_cq_for_docx(cq, doc, is_math=False):
    """Formats and writes a single CQ and its solution to the output document."""
    q_labels = QUESTION_LABELS[:-1] if is_math else QUESTION_LABELS

    # 1. Stem and Questions
    # Render the stem with the first line bold
    for i, para in enumerate(cq.get('stem_meta', [])):
        is_bold = (i == 0)
        render_parts_to_para(doc.add_paragraph(), split_text_and_omml(para), bold=is_bold)

    # Render each question, bolded
    for label in q_labels:
        if label in cq.get('questions_meta', {}):
            para = cq['questions_meta'][label][0]
            render_parts_to_para(doc.add_paragraph(), split_text_and_omml(para), bold=True)
    
    # 2. Solution Header
    p_header = doc.add_paragraph()
    run = p_header.add_run(f"{cq['serial']} নং প্রশ্নের সমাধান")
    set_bangla_font(run, bold=True, color=SOLUTION_HEADER_COLOR, underline=True)
    
    # 3. Solutions
    for label in q_labels:
        if label in cq.get('solutions_meta', {}):
            solution_paras = cq['solutions_meta'][label]
            
            # Start the paragraph for the solution part
            p_sol = doc.add_paragraph()
            
            # Manually add the bolded label (e.g., "ক. ")
            set_bangla_font(p_sol.add_run(f"{label}. "), bold=True)
            
            # --- Render First Paragraph of Solution ---
            first_para_parts = split_text_and_omml(solution_paras[0])
            prefix_regex = fr'^\s*উত্তর\s*[\(]?\s*{label}\s*[\.\)]?\s*\.?\s*'
            
            # Use the new helper function to robustly strip the prefix
            cleaned_parts = strip_prefix_from_parts(first_para_parts, prefix_regex)
            
            # Render the cleaned parts into the same paragraph (without extra bolding)
            render_parts_to_para(p_sol, cleaned_parts)

            # --- Render Subsequent Paragraphs of the same answer ---
            for para in solution_paras[1:]:
                render_parts_to_para(doc.add_paragraph(), split_text_and_omml(para))
    
    # Add space after the entire CQ block for better separation
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

def convert_file(src_file, out_file, is_math):
    try:
        doc = Document(src_file)
    except Exception as e:
        messagebox.showerror("File Error", f"Could not open source file.\n\n{e}")
        return False, None

    cqs = extract_cqs(doc.paragraphs, is_math)
    if not cqs:
        messagebox.showwarning("No CQs Found", "Could not find any CQs with the specified structure in the document.")
        return False, None
        
    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(10.65)
    section.left_margin = Inches(0.8); section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5); section.bottom_margin = Inches(0.3)
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2'); cols.set(qn('w:space'), '210')

    for cq in cqs:
        format_cq_for_docx(cq, outdoc, is_math)
        
    try:
        outdoc.save(out_file)
    except Exception as e:
        messagebox.showerror("Save Error", f"Could not save output DOCX file.\n\n{e}")
        return False, None

    return True, cqs

# --- Tkinter GUI ---
def main_gui():
    def select_input():
        filename = filedialog.askopenfilename(title="Select Source DOCX file", filetypes=[("Word Files", "*.docx")])
        if filename: input_entry.delete(0, tk.END); input_entry.insert(0, filename)

    def select_output():
        filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")], title="Save Formatted DOCX as")
        if filename: output_entry.delete(0, tk.END); output_entry.insert(0, filename)

    def run_conversion():
        src_file = input_entry.get().strip()
        out_file = output_entry.get().strip()
        is_math_subject = var_is_math.get()

        if not src_file or not out_file:
            messagebox.showwarning("Input Needed", "Please select both input and output files!")
            return
        
        root.config(cursor="wait"); root.update()
        try:
            ok, cqs_data = convert_file(src_file, out_file, is_math_subject)
            if ok and cqs_data:
                basepath = out_file.rsplit('.', 1)[0]
                rows = cqs_to_rows(cqs_data, is_math_subject)
                df = pd.DataFrame(rows)
                if var_excel.get(): df.to_excel(basepath + ".xlsx", index=False)
                if var_csv.get(): df.to_csv(basepath + ".csv", index=False, encoding='utf-8-sig')
                if var_json.get(): df.to_json(basepath + ".json", orient='records', indent=2, force_ascii=False)
                messagebox.showinfo("Success!", f"Successfully converted and exported CQ data!\n\nFiles saved with base name: {basepath}")
        finally:
            root.config(cursor="")

    root = tk.Tk()
    root.title("CQ Sheet Formatter & Exporter")

    tk.Label(root, text="Source DOCX:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
    input_entry = tk.Entry(root, width=60)
    input_entry.grid(row=0, column=1, padx=5, columnspan=2)
    tk.Button(root, text="Browse...", command=select_input).grid(row=0, column=3, padx=5)

    tk.Label(root, text="Output DOCX:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
    output_entry = tk.Entry(root, width=60)
    output_entry.grid(row=1, column=1, padx=5, columnspan=2)
    tk.Button(root, text="Browse...", command=select_output).grid(row=1, column=3, padx=5)
    
    options_frame = tk.LabelFrame(root, text="Options", padx=10, pady=10)
    options_frame.grid(row=2, column=0, columnspan=4, padx=10, pady=10, sticky="ew")

    var_is_math = tk.BooleanVar(value=False)
    tk.Checkbutton(options_frame, text="Math Subject (No 'ঘ' Question)", variable=var_is_math).grid(row=0, column=0, sticky="w")
    
    tk.Label(options_frame, text="Also Export As:").grid(row=1, column=0, pady=(10,0), sticky="w")
    var_excel = tk.IntVar(value=1)
    var_csv = tk.IntVar(value=1)
    var_json = tk.IntVar(value=0)
    tk.Checkbutton(options_frame, text="Excel (.xlsx)", variable=var_excel).grid(row=2, column=0, sticky="w")
    tk.Checkbutton(options_frame, text="CSV (.csv)", variable=var_csv).grid(row=2, column=1, sticky="w")
    tk.Checkbutton(options_frame, text="JSON (.json)", variable=var_json).grid(row=2, column=2, sticky="w")

    tk.Button(root, text="Process CQ File", command=run_conversion, bg="#13825c", fg="white", width=20, height=2, font=("Arial", 10, "bold")).grid(row=3, column=1, columnspan=2, pady=15)

    root.mainloop()

if __name__ == "__main__":
    main_gui()