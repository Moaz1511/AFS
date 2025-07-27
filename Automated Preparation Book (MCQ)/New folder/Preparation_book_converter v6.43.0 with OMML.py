from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches
import re
from lxml import etree

def patch_omml_font_size(omml_xml, size_pt=11):
    size_val = str(int(size_pt * 2))
    tree = etree.fromstring(omml_xml)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for sz_tag in tree.xpath('.//w:sz', namespaces=ns):
        sz_tag.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = size_val
    for sz_tag in tree.xpath('.//w:szCs', namespaces=ns):
        sz_tag.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = size_val
    return etree.tostring(tree, encoding='unicode')

def set_tiro_bangla_font(run, font_size=11):
    run.font.name = "Tiro Bangla"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
    run.font.size = Pt(font_size)

def add_text_and_equations(para, text, ommls, font_size=11):
    if text:
        run = para.add_run(text)
        set_tiro_bangla_font(run, font_size)
    if ommls:
        for omml_xml in ommls:
            omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=font_size)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_xml_patched +
                '</w:r>'
            )
            para._p.append(omml_run)

def extract_option_parts(src_para, label):
    """Extract both text and OMML for a single option"""
    para_xml = etree.tostring(src_para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    textval = ""
    ommls = []
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            texts = node.findall('.//w:t', namespaces=ns)
            for t in texts:
                txt = t.text or ""
                if txt.strip().startswith(f"{label}."):
                    txt = txt.strip()[len(f"{label}."):].lstrip()
                textval += txt
        elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
            ommls.append(etree.tostring(node, encoding='unicode'))
    return textval.strip(), ommls

def split_text_and_equations(src_para, label=None):
    """
    Extracts the text (with label optionally removed) and OMML math parts from a paragraph.
    Returns: (text_string, [omml_xml, ...])
    """
    from lxml import etree
    para_xml = etree.tostring(src_para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    }
    textval = ""
    ommls = []
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            texts = node.findall('.//w:t', namespaces=ns)
            for t in texts:
                txt = t.text or ""
                if label and txt.strip().startswith(f"{label}."):
                    txt = txt.strip()[len(f"{label}."):].lstrip()
                textval += txt
        elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
            ommls.append(etree.tostring(node, encoding='unicode'))
    return textval.strip(), ommls

def write_answer_line(doc, mcq):
    ans_p = doc.add_paragraph()
    # 1. "উত্তর:" colored
    run = ans_p.add_run("উত্তর:")
    run.font.name = "Tiro Bangla"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x08, 0x85, 0x65)
    
    ans_label = mcq.get('answer', '').strip()
    answer_option = mcq['options'].get(ans_label, None)

    # 2. Add label with color
    run2 = ans_p.add_run(f" {ans_label}. ")
    run2.font.name = "Tiro Bangla"
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run2.font.size = Pt(11)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0x08, 0x85, 0x65)  # color the label too

    # 3. Show answer text/equation (if any)
    if answer_option:
        opt_text, para_or_omml = answer_option
        try:
            # Split into text/equation
            textval, ommls = split_text_and_equations(para_or_omml, ans_label)
            # Add text part (if any)
            if textval:
                run3 = ans_p.add_run(textval)
                run3.font.name = "Tiro Bangla"
                run3.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
                run3.font.size = Pt(11)
            # Add OMML equations (if any)
            for omml_xml in ommls:
                omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=11)
                omml_run = parse_xml(
                    '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                    'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                    omml_xml_patched +
                    '</w:r>'
                )
                ans_p._p.append(omml_run)
        except Exception as e:
            # fallback: just text
            run3 = ans_p.add_run(opt_text)
            run3.font.name = "Tiro Bangla"
            run3.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
            run3.font.size = Pt(11)
    else:
        # fallback if no answer option, use MCQ stored answer text
        ans_text = mcq.get('answer_text', '').strip()
        if ans_text:
            run3 = ans_p.add_run(ans_text)
            run3.font.name = "Tiro Bangla"
            run3.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
            run3.font.size = Pt(11)

    # Set spacing after answer line
    ans_p.paragraph_format.space_after = Pt(0.6)
    ans_p.paragraph_format.space_before = Pt(0)




def format_mcq(mcq, doc):
    qline0, style0, src_para0 = mcq['question_lines_meta'][0]
    serial = mcq['serial']
    labels = ['ক', 'খ', 'গ', 'ঘ']

    # --- Question + Reference ---
    q_para = doc.add_paragraph()
    run = q_para.add_run(f"{serial}. ")
    set_tiro_bangla_font(run)
    para_xml = etree.tostring(src_para0._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            texts = node.findall('.//w:t', namespaces=ns)
            for t in texts:
                text = t.text
                if text and text.strip().startswith(f"{serial}."):
                    text = text.strip()[len(f"{serial}."):].lstrip()
                run_q = q_para.add_run(text)
                set_tiro_bangla_font(run_q)
        elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
            omml_xml = etree.tostring(node, encoding='unicode')
            omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=11)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_xml_patched +
                '</w:r>'
            )
            q_para._p.append(omml_run)
    if mcq.get('reference'):
        run_ref = q_para.add_run(f" {mcq['reference']}")
        set_tiro_bangla_font(run_ref)

    # Question block spacing
    q_para.paragraph_format.space_after = Pt(0)
    q_para.paragraph_format.space_before = Pt(0)

    # List-style sublines (roman) or extra text
    list_index = 1
    roman = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']
    for qline, style, src_para in mcq['question_lines_meta'][1:]:
        if 'List' in style:
            prefix = roman[list_index-1] if list_index <= len(roman) else f"{list_index}."
            p = doc.add_paragraph(prefix + " ")
            copy_paragraph_with_equations_and_style(src_para, p)
            # write_answer_line(doc, mcq)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            list_index += 1
        elif qline.strip():
            p = doc.add_paragraph()
            copy_paragraph_with_equations_and_style(src_para, p)
            # write_answer_line(doc, mcq)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

    # --- Options: robust per-option handling ---
    opts = []
    maxlen = 0
    for k in labels:
        v = mcq['options'].get(k)
        if v and v[1]:
            opt_text, opt_ommls = extract_option_parts(v[1], k)
        else:
            opt_text, opt_ommls = v[0] if v else "", []
        opts.append((k, opt_text, opt_ommls))
        maxlen = max(maxlen, len(opt_text))
    LONG_LIMIT = 22  # Adjust this for what you want as "too long"

    # Decide on single vs two-per-line
    all_option_is_equation = all(opt[1] == "" and opt[2] for opt in opts)
    if maxlen > LONG_LIMIT or all_option_is_equation:
        # One per line
        for idx, (label, opt_text, ommls) in enumerate(opts):
            para = doc.add_paragraph()
            run_label = para.add_run(f"{label}. ")
            set_tiro_bangla_font(run_label)
            if opt_text:
                run_opt = para.add_run(opt_text)
                set_tiro_bangla_font(run_opt)
            if ommls:
                para.add_run(" ")
                add_text_and_equations(para, "", ommls)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
    else:
        # Two options per line
        # First row: ক, খ
        opt_line1 = doc.add_paragraph()
        for idx in [0, 1]:
            label, opt_text, ommls = opts[idx]
            run_label = opt_line1.add_run(f"{label}. ")
            set_tiro_bangla_font(run_label)
            if opt_text:
                run_opt = opt_line1.add_run(opt_text)
                set_tiro_bangla_font(run_opt)
            if ommls:
                opt_line1.add_run(" ")
                add_text_and_equations(opt_line1, "", ommls)
            opt_line1.add_run("\t\t")
        opt_line1.paragraph_format.space_after = Pt(0)
        opt_line1.paragraph_format.space_before = Pt(0)
        # Second row: গ, ঘ
        opt_line2 = doc.add_paragraph()
        for idx in [2, 3]:
            label, opt_text, ommls = opts[idx]
            run_label = opt_line2.add_run(f"{label}. ")
            set_tiro_bangla_font(run_label)
            if opt_text:
                run_opt = opt_line2.add_run(opt_text)
                set_tiro_bangla_font(run_opt)
            if ommls:
                opt_line2.add_run(" ")
                add_text_and_equations(opt_line2, "", ommls)
            opt_line2.add_run("\t\t")
        opt_line2.paragraph_format.space_after = Pt(0)
        opt_line2.paragraph_format.space_before = Pt(0)

    # --- Answer ---
    write_answer_line(doc, mcq)
    # ans = mcq.get('answer', '')
    # ans_text = mcq.get('answer_text', '')
    # ans_para = doc.add_paragraph()
    # run = ans_para.add_run("উত্তর:")
    # set_tiro_bangla_font(run)
    # run.bold = True
    # run.font.color.rgb = RGBColor(0x08, 0x85, 0x65)
    # run2 = ans_para.add_run(f" {ans}. {ans_text}")
    # set_tiro_bangla_font(run2)
    # # If correct option is an equation, show the equation too!
    # if ans in ['ক','খ','গ','ঘ']:
    #     ans_idx = labels.index(ans)
    #     if opts[ans_idx][2]:
    #         ans_para.add_run(" ")
    #         add_text_and_equations(ans_para, "", opts[ans_idx][2])
    # # Spacing after answer (between questions)
    # ans_para.paragraph_format.space_before = Pt(0)
    # ans_para.paragraph_format.space_after = Pt(7.2)  # ~0.6 line
    

def copy_paragraph_with_equations_and_style(src_para, out_para, font_name="Tiro Bangla", font_size=11, color_hex=None, bold=False, color_answer_label_only=False):
    from lxml import etree
    para_xml = etree.tostring(src_para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            texts = node.findall('.//w:t', namespaces=ns)
            for t in texts:
                text = t.text or ""
                # ---- If line starts with "উত্তর:" and special flag is set ----
                if color_answer_label_only and text.strip().startswith("উত্তর:"):
                    idx = text.find("উত্তর:")
                    # Add "উত্তর:" with color
                    run1 = out_para.add_run(text[idx:idx+6])
                    run1.font.name = font_name
                    run1.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run1.font.size = Pt(font_size)
                    run1.bold = True
                    run1.font.color.rgb = RGBColor(0x08, 0x85, 0x65)  # or use color_hex
                    # Add rest with normal color
                    rest = text[idx+6:]
                    if rest:
                        run2 = out_para.add_run(rest)
                        run2.font.name = font_name
                        run2.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        run2.font.size = Pt(font_size)
                        run2.bold = bold
                    continue
                # ---- Normal case: ----
                run = out_para.add_run(text)
                run.font.name = font_name
                run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = Pt(font_size)
                if color_hex:
                    run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))
                run.font.bold = bold
        elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
            omml_xml = etree.tostring(node, encoding='unicode')
            omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=font_size)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_xml_patched +
                '</w:r>'
            )
            out_para._p.append(omml_run)
    p_format = out_para.paragraph_format
    p_format.line_spacing = 1.3
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    return out_para


# def split_text_and_equations(src_para):
#     """
#     Extracts the text content and OMML (math) parts from a paragraph.
#     Returns: (text_string, [omml_xml_1, omml_xml_2, ...])
#     """
#     from lxml import etree
#     para_xml = etree.tostring(src_para._element, encoding='unicode')
#     tree = etree.fromstring(para_xml.encode('utf-8'))
#     ns = {
#         'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
#         'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
#     }
#     textval = ""
#     ommls = []
#     for node in tree.iterchildren():
#         if node.tag.endswith('r'):
#             texts = node.findall('.//w:t', namespaces=ns)
#             for t in texts:
#                 textval += t.text or ""
#         elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
#             ommls.append(etree.tostring(node, encoding='unicode'))
#     return textval.strip(), ommls


def extract_mcqs(paragraphs):
    mcqs = []
    q_re = re.compile(r'^(\d+)\.\s*(.*)')
    ref_re = re.compile(r'^\[(.+)\]')
    opt_re = re.compile(r'^(ক|খ|গ|ঘ)\.\s*(.+)')
    ans_re = re.compile(r'^উত্তর[:：]\s*(\w)\.?\s*(.*)')
    state = 0
    cur = {}
    for para in paragraphs:
        text = para.text.strip()
        style = para.style.name if hasattr(para, "style") else ""
        if not text:
            continue
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if state == 0:
                m = q_re.match(line)
                if m:
                    cur = {
                        'serial': m.group(1),
                        'question_lines_meta': [(m.group(2), style, para)],
                        'reference': '',
                        'options': {},
                        'answer': ''
                    }
                    state = 1
            elif state == 1:
                m = ref_re.match(line)
                if m:
                    cur['reference'] = m.group(0)
                    continue
                if opt_re.match(line):
                    state = 2
                else:
                    cur['question_lines_meta'].append((line, style, para))
                    continue
            if state == 2:
                m = opt_re.match(line)
                if m:
                    cur['options'][m.group(1)] = (m.group(2), para)
                    continue
                elif ans_re.match(line):
                    m = ans_re.match(line)
                    cur['answer'] = m.group(1)
                    cur['answer_text'] = cur['options'].get(cur['answer'], ('', ''))[0]
                    mcqs.append(cur)
                    state = 0
                else:
                    if cur['options']:
                        last_opt = list(cur['options'])[-1]
                        val, oxmls = cur['options'][last_opt]
                        cur['options'][last_opt] = (val + ' ' + line, para)
    return mcqs

def main():
    src = "AP_-_MCQ_Sheet_-_Class_6_-_Chapter_1.1^J_1.2^J_1.3^J_1.4^J_1.5^J_1.6_-_স্বাভাবিক_সংখ্যা_ও_ভগ্নাংশ.docx"
    out = "Reformatted_MCQ_Sheet.docx"
    doc = Document(src)

    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.6)
    section.footer_distance = Inches(0.2)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # Set to 2 columns

    mcqs = extract_mcqs(doc.paragraphs)
    for mcq in mcqs:
        format_mcq(mcq, outdoc)

    outdoc.save(out)
    print("Done! Output:", out)

if __name__ == "__main__":
    main()
