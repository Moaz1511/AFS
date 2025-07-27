import re
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches

FONT_NAME = "Tiro Bangla"
FONT_SIZE = 11
ANSWER_COLOR = RGBColor(0x08, 0x85, 0x65)
OPTION_LABELS = ['ক', 'খ', 'গ', 'ঘ']
LONG_OPTION_LIMIT = 15  # character length

def patch_omml_font_size(omml_xml, size_pt=FONT_SIZE):
    size_val = str(int(size_pt * 2))
    tree = etree.fromstring(omml_xml)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for sz_tag in tree.xpath('.//w:sz', namespaces=ns):
        sz_tag.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = size_val
    for sz_tag in tree.xpath('.//w:szCs', namespaces=ns):
        sz_tag.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = size_val
    return etree.tostring(tree, encoding='unicode')

def set_bangla_font(run, font_size=FONT_SIZE, bold=False, color=None):
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def clean_punct_spacing(text):
    # Remove spaces before , and ?
    # Works for both Bengali and English , ?
    text = re.sub(r'\s+([,?\]])', r'\1', text)
    text = re.sub(r'(\[)\s+', r'\1', text)
    # text = re.sub(r'\s+([,?])', r'\1', text)
    # Also handle Bengali comma (।) if you want: r'\s+(।)', r'।'
    return text

def clean_all_bracket_spacing(text):
    text = re.sub(r'([\[\(\{])\s+', r'\1', text)
    text = re.sub(r'\s+([\]\)\}])', r'\1', text)
    return text

def split_text_and_omml(para, strip_label=None):
    para_xml = etree.tostring(para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    parts = []
    label_stripped = False
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            texts = node.findall('.//w:t', namespaces=ns)
            for t in texts:
                txt = t.text or ""
                if strip_label and not label_stripped and txt.strip().startswith(f"{strip_label}."):
                    txt = txt.strip()[len(f"{strip_label}."):].lstrip()
                    label_stripped = True
                # --- Key Fix: always add a space between runs unless both are punctuation ---
                if txt.strip():
                    if parts and parts[-1][0] == "text" and not txt.startswith(('।', '.', ',', '?', '!', '’', '”')):
                        # Insert a space unless the new part is punctuation
                        parts.append(("text", ' '))
                    parts.append(("text", txt))
        elif node.tag.endswith(('oMath', 'oMathPara')):
            omml_xml = etree.tostring(node, encoding='unicode')
            parts.append(("omml", omml_xml))
    return parts

def extract_mcqs(paragraphs):
    q_re = re.compile(r'^\s*(\d+)\.\s*(.*)')
    ref_re = re.compile(r'\s*\[(.+)\]\s*')
    opt_re = re.compile(r'^\s*(ক|খ|গ|ঘ)\.\s*(.*)')
    ans_re = re.compile(r'^\s*উত্তর[:：]\s*(\w)\.?\s*(.*)')

    mcqs = []
    cur = None
    state = 0

    for para in paragraphs:
        text = para.text.strip() if para.text else ""
        if not text:
            continue
        if state == 0:
            m = q_re.match(text)
            if m:
                cur = {
                    'serial': m.group(1),
                    'question_meta': [(m.group(2), para)],
                    'reference': '',
                    'options_meta': {},
                    'answer_label': '',
                }
                state = 1
        elif state == 1:
            mref = ref_re.match(text)
            mopt = opt_re.match(text)
            if mref:
                cur['reference'] = mref.group(0).strip()
            elif mopt:
                cur['options_meta'][mopt.group(1)] = (mopt.group(2), para)
                state = 2
            else:
                cur['question_meta'].append((text, para))
        elif state == 2:
            mans = ans_re.match(text)
            mopt = opt_re.match(text)
            if mopt:
                cur['options_meta'][mopt.group(1)] = (mopt.group(2), para)
            elif mans:
                cur['answer_label'] = mans.group(1)
                mcqs.append(cur)
                state = 0
            else:
                if cur['options_meta']:
                    last_label = list(cur['options_meta'])[-1]
                    prev_text, prev_para = cur['options_meta'][last_label]
                    cur['options_meta'][last_label] = (f"{prev_text} {text}", prev_para)
    return mcqs

def render_parts_to_para(para, parts):
    last_type = None
    last_text = ''
    for idx, (ctype, cval) in enumerate(parts):
        # Determine if a space should be added
        need_space = False
        if idx > 0:
            prev_type, prev_val = parts[idx - 1]
            if (
                (prev_type == "omml" and ctype == "text" and cval and not cval[0] in "।.,?!”’") or
                (prev_type == "text" and prev_val and not prev_val[-1] in "।.,?!”’" and ctype == "omml")
            ):
                need_space = True
        # Add space if needed
        if need_space:
            para.add_run(" ")
        # Render as before
        if ctype == "text":
            runopt = para.add_run(clean_punct_spacing(cval))
            set_bangla_font(runopt)
        elif ctype == "omml":
            omml_xml_patched = patch_omml_font_size(cval)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_xml_patched +
                '</w:r>'
            )
            para._p.append(omml_run)
        last_type = ctype

def write_colored_answer_line(doc, ans_label, ans_tuple):
    para = doc.add_paragraph()
    run1 = para.add_run("উত্তর:")
    set_bangla_font(run1, bold=True, color=ANSWER_COLOR)
    run2 = para.add_run(f" {ans_label}. ")
    set_bangla_font(run2, bold=False)
    if ans_tuple:
        parts = split_text_and_omml(ans_tuple[1], strip_label=ans_label)
        render_parts_to_para(para, parts)
    para.paragraph_format.space_after = Pt(7)
    para.paragraph_format.space_before = Pt(0)

def format_mcq(mcq, doc):
    # --- Question and reference ---
    q_para = doc.add_paragraph()
    run = q_para.add_run(f"{mcq['serial']}. ")
    set_bangla_font(run, bold=True)
    
    # --- Only show the serial once, not duplicated ---
    first_qtext, first_qpara = mcq['question_meta'][0]
    parts = split_text_and_omml(first_qpara)
    if parts and parts[0][0] == "text":
        # Remove serial number (e.g., "১. ") if present in the first part
        parts[0] = ("text", re.sub(r'^\s*\d+\.\s*', '', parts[0][1], count=1))
    render_parts_to_para(q_para, parts)
    
    # Add reference if present (clean space before and after [ ])
    if mcq.get('reference'):
        ref_clean = clean_all_bracket_spacing(clean_punct_spacing(mcq['reference']))
        run_ref = q_para.add_run(" " + ref_clean)
        set_bangla_font(run_ref)

    q_para.paragraph_format.space_after = Pt(0)
    q_para.paragraph_format.space_before = Pt(0)

    # --- Patch: Detect subparts and prompt ---
    subparts = []
    prompt_line = None
    for text, para in mcq['question_meta'][1:]:
        # Detect if line is a prompt/question about the subparts (flexible for any phrase)
        if any(key in text for key in ["সঠিক", "ঠিক", "যথাযথ"]):
            prompt_line = (text, para)
        else:
            subparts.append((text, para))

    # Render subparts with roman numerals, only if present
    roman = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']
    for idx, (qtext, qpara) in enumerate(subparts):
        sub_para = doc.add_paragraph()
        run_r = sub_para.add_run(f"{roman[idx]} ")
        set_bangla_font(run_r, bold=True)
        render_parts_to_para(sub_para, split_text_and_omml(qpara))
        sub_para.paragraph_format.space_after = Pt(0)
        sub_para.paragraph_format.space_before = Pt(0)

    # Patch: Print the prompt line if present
    if prompt_line:
        cpara = doc.add_paragraph()
        runc = cpara.add_run(prompt_line[0])
        set_bangla_font(runc)
        cpara.paragraph_format.space_after = Pt(0)
        cpara.paragraph_format.space_before = Pt(0)

    # --- Option layout (keep your original logic, but fix space bug) ---
    opts = []
    maxlen = 0
    for label in OPTION_LABELS:
        opt_tuple = mcq['options_meta'].get(label)
        textlen = 0
        if opt_tuple:
            # get the full cleaned text length
            for ctype, cval in split_text_and_omml(opt_tuple[1], strip_label=label):
                if ctype == "text":
                    textlen += len(cval)
        maxlen = max(maxlen, textlen)
        opts.append((label, opt_tuple))

    def is_option_long(opt_tuple, label):
        textlen = 0
        has_omml = False
        if not opt_tuple:
            return False
        for ctype, cval in split_text_and_omml(opt_tuple[1], strip_label=label):
            if ctype == "text":
                textlen += len(cval)
            elif ctype == "omml":
                has_omml = True
        return textlen > LONG_OPTION_LIMIT or has_omml

    rows = [(0, 1), (2, 3)]
    for r0, r1 in rows:
        long0 = is_option_long(opts[r0][1], opts[r0][0])
        long1 = is_option_long(opts[r1][1], opts[r1][0])
        # If either option is long, print as single
        if long0:
            para = doc.add_paragraph()
            label, opt_tuple = opts[r0]
            run_lab = para.add_run(f"{label}. ")
            set_bangla_font(run_lab, bold=True)
            if opt_tuple:
                render_parts_to_para(para, split_text_and_omml(opt_tuple[1], strip_label=label))
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
        if long1:
            para = doc.add_paragraph()
            label, opt_tuple = opts[r1]
            run_lab = para.add_run(f"{label}. ")
            set_bangla_font(run_lab, bold=True)
            if opt_tuple:
                render_parts_to_para(para, split_text_and_omml(opt_tuple[1], strip_label=label))
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
        # Otherwise, print both as a single 2-column line
        if not long0 and not long1:
            para = doc.add_paragraph()
            for idx in [r0, r1]:
                label, opt_tuple = opts[idx]
                run_lab = para.add_run(f"{label}. ")
                set_bangla_font(run_lab, bold=True)
                if opt_tuple:
                    for ctype, cval in split_text_and_omml(opt_tuple[1], strip_label=label):
                        if ctype == "text":
                            runopt = para.add_run(clean_punct_spacing(cval))  # <-- PATCH: do NOT .strip()
                            set_bangla_font(runopt)
                        elif ctype == "omml":
                            omml_xml_patched = patch_omml_font_size(cval)
                            omml_run = parse_xml(
                                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                                omml_xml_patched +
                                '</w:r>'
                            )
                            para._p.append(omml_run)
                para.add_run('\t\t')
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)

    # --- Answer line ---
    ans_label = mcq.get('answer_label')
    ans_tuple = mcq['options_meta'].get(ans_label)
    write_colored_answer_line(doc, ans_label, ans_tuple)


def main():
    src = "AP_-_MCQ_Sheet_-_Class_6_-_Chapter_1.1^J_1.2^J_1.3^J_1.4^J_1.5^J_1.6_-_স্বাভাবিক_সংখ্যা_ও_ভগ্নাংশ.docx"
    out = "Reformatted_MCQ_Sheet.docx"
    doc = Document(src)
    outdoc = Document()
    # Layout: 2-column
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    mcqs = extract_mcqs(doc.paragraphs)
    for mcq in mcqs:
        format_mcq(mcq, outdoc)
    outdoc.save(out)
    print("Done! Output:", out)

if __name__ == "__main__":
    main()
