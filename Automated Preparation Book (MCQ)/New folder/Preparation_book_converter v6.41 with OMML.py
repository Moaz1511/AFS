from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches
import re
from lxml import etree

def get_omml_xml(para):
    return [etree.tostring(omml, encoding="unicode") for omml in para._element.xpath('.//m:oMath | .//m:oMathPara')]


def patch_omml_font_size(omml_xml, size_pt=11):
    from lxml import etree
    size_val = str(int(size_pt * 2))
    tree = etree.fromstring(omml_xml)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for sz_tag in tree.xpath('.//w:sz', namespaces=ns):
        sz_tag.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = size_val
    for sz_tag in tree.xpath('.//w:szCs', namespaces=ns):
        sz_tag.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = size_val
    return etree.tostring(tree, encoding='unicode')

def add_text_and_equations(para, text, ommls, font_name="Tiro Bangla", font_size=11):
    run = para.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    if ommls:
        for omml_xml in ommls:
            omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=font_size)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_xml_patched +
                '</w:r>'
            )
            para._p.append(omml_run)

def set_tiro_bangla_font(para):
    for run in para.runs:
        run.font.name = "Tiro Bangla"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
        run.font.size = Pt(11)
    # Set line spacing and paragraph spacing
    p_format = para.paragraph_format
    p_format.line_spacing = 1.3
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

def copy_paragraph_with_equations_and_style(src_para, out_para, font_name="Tiro Bangla", font_size=11, color_hex=None, bold=False):
    from lxml import etree
    para_xml = etree.tostring(src_para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            texts = node.findall('.//w:t', namespaces=ns)
            for t in texts:
                run = out_para.add_run(t.text)
                run.font.name = font_name
                run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = Pt(font_size)
                if color_hex:
                    run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))
                run.font.bold = bold
        elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
            omml_xml = etree.tostring(node, encoding='unicode')
            omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=font_size)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_xml_patched +
                '</w:r>'
            )
            out_para._p.append(omml_run)
    p_format = out_para.paragraph_format
    p_format.line_spacing = 1.3
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    return out_para

def format_mcq(mcq, doc):
    # Question line (serial only once)
    q_para = doc.add_paragraph()
    serial = mcq['serial']
    run = q_para.add_run(f"{serial}. ")
    run.font.name = "Tiro Bangla"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
    run.font.size = Pt(11)
    qline, _, ommls = mcq['question_lines_meta'][0]
    add_text_and_equations(q_para, qline, ommls)

    # Additional lines (lists, multiline question)
    list_index = 1
    roman = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']
    for qline, style, ommls in mcq['question_lines_meta'][1:]:
        if 'List' in style:
            prefix = roman[list_index-1] if list_index <= len(roman) else f"{list_index}."
            p = doc.add_paragraph(prefix + " ")
            add_text_and_equations(p, qline, ommls)
            list_index += 1
        elif qline.strip():
            p = doc.add_paragraph()
            add_text_and_equations(p, qline, ommls)

    # Reference
    if mcq.get('reference'):
        ref_para = doc.add_paragraph(mcq['reference'])
        set_tiro_bangla_font(ref_para)

    # Options (two per line)
    opts = []
    for k in ['ক', 'খ', 'গ', 'ঘ']:
        v = mcq['options'].get(k, ('', []))
        opt_text, ommls = v
        opts.append((k, opt_text, ommls))
    # First line: ক, খ
    opt_line1 = doc.add_paragraph()
    for idx in [0, 1]:
        label, opt_text, ommls = opts[idx]
        run = opt_line1.add_run(f"{label}. ")
        run.font.name = "Tiro Bangla"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
        run.font.size = Pt(11)
        add_text_and_equations(opt_line1, opt_text, ommls)
        opt_line1.add_run("\t\t")
    # Second line: গ, ঘ
    opt_line2 = doc.add_paragraph()
    for idx in [2, 3]:
        label, opt_text, ommls = opts[idx]
        run = opt_line2.add_run(f"{label}. ")
        run.font.name = "Tiro Bangla"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
        run.font.size = Pt(11)
        add_text_and_equations(opt_line2, opt_text, ommls)
        opt_line2.add_run("\t\t")

    # Answer line (only "উত্তর:" colored)
    ans = mcq.get('answer', '')
    ans_text = mcq.get('answer_text', '')
    ans_para = doc.add_paragraph()
    run = ans_para.add_run("উত্তর:")
    run.font.name = "Tiro Bangla"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0x08, 0x85, 0x65)
    run2 = ans_para.add_run(f" {ans}. {ans_text}")
    run2.font.name = "Tiro Bangla"
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
    run2.font.size = Pt(11)
    # doc.add_paragraph('')



def extract_mcqs(paragraphs):
    mcqs = []
    q_re = re.compile(r'^(\d+)\.\s*(.*)')
    ref_re = re.compile(r'^\[(.+)\]')
    opt_re = re.compile(r'^(ক|খ|গ|ঘ)\.\s*(.+)')
    ans_re = re.compile(r'^উত্তর[:：ঃ]\s*(\w)\.?\s*(.*)')  # Added 'ঃ'
    state = 0
    cur = {}
    for para in paragraphs:
        text = para.text.strip()
        style = para.style.name if hasattr(para, "style") else ""
        ommls = get_omml_xml(para)
        if not text and ommls:
            text = "[Equation]"
        if not text:
            continue

        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if state == 0:
                m = q_re.match(line)
                if m:
                    cur = {
                        'serial': m.group(1),
                        'question_lines_meta': [(m.group(2), style, ommls)],
                        'reference': '',
                        'options': {},
                        'answer': ''
                    }
                    state = 1
            elif state == 1:
                m_ref = ref_re.match(line)
                m_opt = opt_re.match(line)
                m_ans = ans_re.match(line)
                if m_ref:
                    cur['reference'] = m_ref.group(0)
                    continue
                elif m_opt:
                    state = 2
                elif m_ans:
                    cur['answer'] = m_ans.group(1)
                    cur['answer_text'] = cur['options'].get(cur['answer'], ('', []))[0]
                    mcqs.append(cur)
                    state = 0
                else:
                    cur['question_lines_meta'].append((line, style, ommls))
                    continue
            if state == 2:
                m_opt = opt_re.match(line)
                m_ans = ans_re.match(line)
                if m_opt:
                    cur['options'][m_opt.group(1)] = (m_opt.group(2), ommls)
                    continue
                elif m_ans:
                    cur['answer'] = m_ans.group(1)
                    cur['answer_text'] = cur['options'].get(cur['answer'], ('', []))[0]
                    mcqs.append(cur)
                    state = 0
                else:
                    if cur['options']:
                        last_opt = list(cur['options'])[-1]
                        val, oxmls = cur['options'][last_opt]
                        cur['options'][last_opt] = (val + ' ' + line, oxmls)
    return mcqs


def main():
    src = "AP_-_MCQ_Sheet_-_Class_6_-_Chapter_1.1^J_1.2^J_1.3^J_1.4^J_1.5^J_1.6_-_স্বাভাবিক_সংখ্যা_ও_ভগ্নাংশ.docx"
    out = "Reformatted_MCQ_Sheet.docx"
    doc = Document(src)

    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.6)
    section.footer_distance = Inches(0.2)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # Set to 2 columns

    mcqs = extract_mcqs(doc.paragraphs)
    for mcq in mcqs:
        format_mcq(mcq, outdoc)

    outdoc.save(out)
    print("Done! Output:", out)

if __name__ == "__main__":
    main()

