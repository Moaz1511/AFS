from docx import Document
from lxml import etree
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
import re

# -------- Add Tkinter imports ----------
import tkinter as tk
from tkinter import filedialog, messagebox
import os


def set_tiro_bangla_font(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Tiro Bangla'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
        run.font.size = Pt(11)
    # Set line spacing and paragraph spacing
    p_format = paragraph.paragraph_format
    p_format.line_spacing = 1.3
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

def color_ans_text(paragraph, color_hex='#088565'):
    # Only color "উত্তর:" in the answer line
    for run in paragraph.runs:
        if run.text.startswith('উত্তর:'):
            run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))
            break

def get_omml_xml(para):
    omml_blocks = para._element.xpath('.//m:oMath | .//m:oMathPara')
    return [etree.tostring(omml, encoding="unicode") for omml in omml_blocks]

def omml_to_unicode(omml_xml):
    """
    Converts *simple* OMML XML to a Unicode string.
    Only supports simple fractions and superscripts/subscripts for demo purposes!
    """
    try:
        tree = etree.fromstring(omml_xml.encode('utf-8'))
        ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}

        # Check for <m:f> (fraction)
        frac = tree.find('.//m:f', namespaces=ns)
        if frac is not None:
            num = frac.find('.//m:num//m:t', namespaces=ns)
            den = frac.find('.//m:den//m:t', namespaces=ns)
            if num is not None and den is not None:
                return f"{num.text}/{den.text}"

        # Check for <m:sup> (superscript)
        sup = tree.find('.//m:sup', namespaces=ns)
        if sup is not None:
            base = sup.find('.//m:e//m:t', namespaces=ns)
            exp = sup.find('.//m:sup//m:t', namespaces=ns)
            if base is not None and exp is not None:
                supers = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹'}
                exp_uni = ''.join(supers.get(c, c) for c in exp.text)
                return f"{base.text}{exp_uni}"

        # Fallback: Try to find all text in <m:t>
        texts = tree.findall('.//m:t', namespaces=ns)
        return ''.join(t.text for t in texts if t is not None)
    except Exception as e:
        return "[Equation]"


def extract_mcqs(paragraphs):
    mcqs = []
    q_re = re.compile(r'^(\d+)\.\s*(.*)')
    ref_re = re.compile(r'^\[(.+)\]')
    opt_re = re.compile(r'^(ক|খ|গ|ঘ)\.\s*(.+)')
    ans_re = re.compile(r'^উত্তর[:：ঃ]\s*(\w)\.?\s*(.*)')  # Added 'ঃ'
    
    state = 0
    cur = {}
    for para in paragraphs:
        text = para.text.strip()
        style = para.style.name if hasattr(para, "style") else ""
        ommls = get_omml_xml(para)
        if not text and ommls:
            text = "[Equation]"  # placeholder
        if not text:
            continue

        # --- Split each paragraph by new line and process each line ---
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if state == 0:
                m = q_re.match(line)
                if m:
                    cur = {
                        'serial': m.group(1),
                        'question_lines_meta': [(m.group(2), style, ommls)],
                        'reference': '',
                        'options': {},
                        'answer': ''
                    }
                    state = 1
            elif state == 1:
                m_ref = ref_re.match(line)
                m_opt = opt_re.match(line)
                m_ans = ans_re.match(line)
                if m_ref:
                    cur['reference'] = m_ref.group(0)
                    continue
                elif m_opt:
                    state = 2
                elif m_ans:
                    cur['answer'] = m_ans.group(1)
                    cur['answer_text'] = cur['options'].get(cur['answer'], ('', []))[0]
                    mcqs.append(cur)
                    state = 0
                else:
                    cur['question_lines_meta'].append((line, style, ommls))
                    continue
            if state == 2:
                m_opt = opt_re.match(line)
                m_ans = ans_re.match(line)
                if m_opt:
                    cur['options'][m_opt.group(1)] = (m_opt.group(2), ommls)
                    continue
                elif m_ans:
                    cur['answer'] = m_ans.group(1)
                    cur['answer_text'] = cur['options'].get(cur['answer'], ('', []))[0]
                    mcqs.append(cur)
                    state = 0
                else:
                    # Multi-line option continuation
                    if cur['options']:
                        last_opt = list(cur['options'])[-1]
                        val, oxmls = cur['options'][last_opt]
                        cur['options'][last_opt] = (val + ' ' + line, oxmls)
    return mcqs

def format_mcq(mcq, doc):
    if not mcq['question_lines_meta']:
        print(f"Warning: MCQ {mcq.get('serial', '?')} has no question lines. Skipping.")
        return

    # Serial + Question
    qline0, style0, ommls0 = mcq['question_lines_meta'][0]
    # Insert Unicode equation if available
    if ommls0:
        unicodes = [omml_to_unicode(oxml) for oxml in ommls0]
        qline0 += " " + " ".join(unicodes)
    q_para = doc.add_paragraph(f"{mcq['serial']}. {qline0}")
    set_tiro_bangla_font(q_para)

    # List
    list_index = 1
    roman = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']
    for qline, style, ommls in mcq['question_lines_meta'][1:]:
        if ommls:
            unicodes = [omml_to_unicode(oxml) for oxml in ommls]
            qline += " " + " ".join(unicodes)
        if 'List' in style:
            prefix = roman[list_index-1] if list_index <= len(roman) else f"{list_index}."
            p = doc.add_paragraph(f"{prefix} {qline}")
            set_tiro_bangla_font(p)
            list_index += 1
        elif qline.strip():
            p = doc.add_paragraph(qline)
            set_tiro_bangla_font(p)
    # Reference
    if mcq.get('reference'):
        ref_para = doc.add_paragraph(f"{mcq['reference']}")
        set_tiro_bangla_font(ref_para)
    # Options (dynamic)
    opts = []
    for k in ['ক','খ','গ','ঘ']:
        val, oxmls = mcq['options'].get(k, ('', []))
        if oxmls:
            val += " " + " ".join([omml_to_unicode(oxml) for oxml in oxmls])
        opts.append(val)
    max_opt_len = max([len(opt) for opt in opts if opt], default=0)
    LONG_LIMIT = 14
    if max_opt_len > LONG_LIMIT:
        for idx, opt in enumerate(['ক','খ','গ','ঘ']):
            p = doc.add_paragraph(f"{opt}. {opts[idx]}")
            set_tiro_bangla_font(p)
    else:
        opt1 = f"ক. {opts[0]}"
        opt2 = f"খ. {opts[1]}"
        opt3 = f"গ. {opts[2]}"
        opt4 = f"ঘ. {opts[3]}"
        p1 = doc.add_paragraph(f"{opt1}\t\t{opt2}")
        set_tiro_bangla_font(p1)
        p2 = doc.add_paragraph(f"{opt3}\t\t{opt4}")
        set_tiro_bangla_font(p2)
    # Answer
    ans = mcq.get('answer', '')
    ans_text = mcq.get('answer_text', '')
    ans_para = doc.add_paragraph()
    run = ans_para.add_run(f"উত্তর:")  # color only this part
    run.font.name = 'Tiro Bangla'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x08, 0x85, 0x65)  # #088565
    run2 = ans_para.add_run(f" {ans}. {ans_text}")
    run2.font.name = 'Tiro Bangla'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run2.font.size = Pt(11)
    # 5. Blank line for separation
    # doc.add_paragraph('')

def main():
    # Initialize Tkinter and hide main window
    root = tk.Tk()
    root.withdraw()

    # Select input file
    src = filedialog.askopenfilename(
        title="Select the source DOCX file",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not src:
        messagebox.showerror("No file selected", "You did not select any input DOCX file.")
        return

    # Suggest output file name based on input
    default_out = os.path.splitext(os.path.basename(src))[0] + "_Reformatted.docx"
    out = filedialog.asksaveasfilename(
        title="Save reformatted DOCX as...",
        initialfile=default_out,
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not out:
        messagebox.showerror("No output file", "You did not select where to save the output file.")
        return

    # Now process as before
    doc = Document(src)
    mcqs = extract_mcqs(doc.paragraphs)

    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.6)
    section.footer_distance = Inches(0.2)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # Set to 2 columns

    for mcq in mcqs:
        format_mcq(mcq, outdoc)

    outdoc.save(out)
    messagebox.showinfo("Done", f"Done!\nOutput: {out}")

if __name__ == "__main__":
    main()
