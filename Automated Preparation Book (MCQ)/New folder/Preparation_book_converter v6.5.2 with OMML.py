import re
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches

# --- Configuration Constants ---
FONT_NAME = "Tiro Bangla"
FONT_SIZE = 11
ANSWER_COLOR = RGBColor(0x08, 0x85, 0x65)
OPTION_LABELS = ['ক', 'খ', 'গ', 'ঘ']
# Option length (character) thresholds for layout selection
SHORT_OPTION_CHAR_LIMIT = 6   # <= all 4 in 1 line
MEDIUM_OPTION_CHAR_LIMIT = 10  # <= 2 lines (2+2)
LONG_OPTION_CHAR_LIMIT = 15    # <= 4 lines (1 per line) # Character length to trigger single-column layout
ROMAN_NUMERALS = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']

# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def clean_text(text):
    if not text:
        return ""
    # Remove space after opening bracket/quote
    text = re.sub(r'([(\[\{«“‘])\s+', r'\1', text)
    # Remove space before closing bracket/quote/punctuation
    text = re.sub(r'\s+([)\]\}»”’.,?।;:])', r'\1', text)
    # Reduce multiple spaces to single (except tabs for columns)
    text = re.sub(r' {2,}', ' ', text)
    return text

def patch_omml_font_size(omml_xml, size_pt=FONT_SIZE):
    try:
        size_val = str(int(size_pt * 2))
        tree = etree.fromstring(omml_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for sz_tag in tree.xpath('.//w:sz|.//w:szCs', namespaces=ns):
            sz_tag.attrib[qn('w:val')] = size_val
        return etree.tostring(tree, encoding='unicode')
    except etree.XMLSyntaxError:
        return omml_xml # Return original on error

def set_bangla_font(run, font_size=FONT_SIZE, bold=False, color=None):
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def split_text_and_omml(para, strip_label_prefix=None):
    para_xml = etree.tostring(para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    parts = []
    first_text_run = True
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            text = "".join(t.text or "" for t in node.findall('.//w:t', namespaces=ns))
            if first_text_run and strip_label_prefix:
                regex = re.compile(r'^\s*' + re.escape(strip_label_prefix) + r'\.\s*')
                if regex.match(text):
                    text = regex.sub('', text, 1)
                    first_text_run = False
            if text:
                parts.append(("text", text))
        elif node.tag.endswith(('oMath', 'oMathPara')):
            omml_xml = etree.tostring(node, encoding='unicode')
            parts.append(("omml", omml_xml))
    return parts

def render_parts_to_para(para, parts, bold=False, color=None):
    for i, (ctype, cvalue) in enumerate(parts):
        # Add space between text and math, if needed
        if i > 0:
            prev_ctype, _ = parts[i-1]
            if prev_ctype != ctype:
                para.add_run(" ")
        if ctype == "text":
            run = para.add_run(clean_text(cvalue))
            set_bangla_font(run, bold=bold, color=color)
        elif ctype == "omml":
            omml_patched = patch_omml_font_size(cvalue)
            omml_run = parse_xml(
                                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                                omml_patched +
                                '</w:r>'
                            )
            para._p.append(omml_run)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)

# For option measurement
def _is_option_long(opt_tuple, label):
    if not opt_tuple:
        return False
    para, _ = opt_tuple
    text_len = 0
    has_omml = False
    for ctype, cval in split_text_and_omml(para, strip_label_prefix=label):
        if ctype == "text":
            text_len += len(cval.strip())
        elif ctype == "omml":
            has_omml = True
    return text_len > LONG_OPTION_CHAR_LIMIT or has_omml

def _all_options_short(options_dict):
    for label in OPTION_LABELS:
        opt_tuple = options_dict.get(label)
        if not opt_tuple:
            return False
        if _is_option_long(opt_tuple, label):
            return False
    return True

def get_option_length_class(options_dict):
    """
    Returns layout type for options:
    - 'oneline': all options are very short (<= SHORT_OPTION_CHAR_LIMIT, no OMML)
    - 'twoline': all options <= MEDIUM_OPTION_CHAR_LIMIT and no OMML
    - 'fourline': if any option > MEDIUM_OPTION_CHAR_LIMIT or has OMML
    """
    found_omml = False
    maxlen = 0
    for label in OPTION_LABELS:
        opt_tuple = options_dict.get(label)
        if not opt_tuple:
            return "fourline"  # If any option missing, best to use fourline (safe fallback)
        para, _ = opt_tuple
        opt_text_len = 0
        for ctype, cval in split_text_and_omml(para, strip_label_prefix=label):
            if ctype == "text":
                opt_text_len += len(cval.strip())
            elif ctype == "omml":
                found_omml = True
        maxlen = max(maxlen, opt_text_len)
    if found_omml or maxlen > MEDIUM_OPTION_CHAR_LIMIT:
        return "fourline"
    elif maxlen > SHORT_OPTION_CHAR_LIMIT:
        return "twoline"
    else:
        return "oneline"

def option_layout_type(options_dict):
    """
    Returns:
        'oneline'   — All short, no OMML.
        'twoline'   — Any OMML (but all options are "short/medium").
        'fourline'  — Any long text or long OMML.
    """
    has_omml = False
    for label in OPTION_LABELS:
        opt_tuple = options_dict.get(label)
        if not opt_tuple:
            return 'fourline'
        para, _ = opt_tuple
        text_len = 0
        omml_in_this = False
        for ctype, cval in split_text_and_omml(para, strip_label_prefix=label):
            if ctype == "text":
                text_len += len(cval.strip())
            elif ctype == "omml":
                omml_in_this = True
        if text_len > LONG_OPTION_CHAR_LIMIT:
            return 'fourline'  # long text triggers 4-line
        if omml_in_this:
            has_omml = True
            # But, if OMML plus very long? Also 4-line.
            if text_len > LONG_OPTION_CHAR_LIMIT:
                return 'fourline'
    if has_omml:
        return 'twoline'
    return 'oneline'

# =============================================================================
# MCQ FORMATTING
# =============================================================================

def format_mcq(mcq, doc, tab_stops, layout):
    # 1. Question and Reference
    q_para = doc.add_paragraph()
    set_bangla_font(q_para.add_run(f"{mcq['serial']}. "), bold=True)
    # Render question and possible sub-lines
    prompt_line = None
    for i, (para, text) in enumerate(mcq['question_meta']):
        if i == 0:
            render_parts_to_para(q_para, split_text_and_omml(para, strip_label_prefix=mcq['serial']))
        else:
            if any(key in text for key in ["সঠিক", "ঠিক", "যথাযথ", " কোনটি", "কোনগুলো"]):
                prompt_line = para
            else:
                sub_para = doc.add_paragraph()
                set_bangla_font(sub_para.add_run(f"{ROMAN_NUMERALS[i-1]} "), bold=True)
                render_parts_to_para(sub_para, split_text_and_omml(para))
    if mcq.get('reference'):
        q_para.add_run(" ")
        set_bangla_font(q_para.add_run(mcq['reference']))
    if prompt_line:
        render_parts_to_para(doc.add_paragraph(), split_text_and_omml(prompt_line))
    # 2. Options
    # option_layout = get_option_length_class(mcq['options_meta'])

    if layout == "oneline":
        # One line for all options with tabs
        option_para = doc.add_paragraph()
        option_para.paragraph_format.tab_stops.clear_all()
        for tab_pos in tab_stops:
            option_para.paragraph_format.tab_stops.add_tab_stop(tab_pos)
        col = 0
        for label in OPTION_LABELS:
            opt_tuple = mcq['options_meta'].get(label)
            if not opt_tuple:
                continue
            if col > 0:
                option_para.add_run('\t')
            run = option_para.add_run(f"{label}. ")
            set_bangla_font(run, bold=True)
            render_parts_to_para(option_para, split_text_and_omml(opt_tuple[0], strip_label_prefix=label))
            col += 1

    elif layout == "twoline":
        # Two lines (2+2), with tabs
        for row in [(0, 1), (2, 3)]:
            option_para = doc.add_paragraph()
            option_para.paragraph_format.tab_stops.clear_all()
            for tab_pos in tab_stops:
                option_para.paragraph_format.tab_stops.add_tab_stop(tab_pos)
            col = 0
            for idx in row:
                label = OPTION_LABELS[idx]
                opt_tuple = mcq['options_meta'].get(label)
                if not opt_tuple:
                    continue
                if col > 0:
                    option_para.add_run('\t')
                run = option_para.add_run(f"{label}. ")
                set_bangla_font(run, bold=True)
                render_parts_to_para(option_para, split_text_and_omml(opt_tuple[0], strip_label_prefix=label))
                col += 1

    else:  # fourline
        for idx in range(4):
            label = OPTION_LABELS[idx]
            opt_tuple = mcq['options_meta'].get(label)
            if not opt_tuple:
                continue
            p = doc.add_paragraph()
            run = p.add_run(f"{label}. ")
            set_bangla_font(run, bold=True)
            render_parts_to_para(p, split_text_and_omml(opt_tuple[0], strip_label_prefix=label))
    # 3. Answer Line
    ans_label = mcq.get('answer_label')
    if ans_label:
        ans_tuple = mcq['options_meta'].get(ans_label)
        p = doc.add_paragraph()
        set_bangla_font(p.add_run("উত্তর: "), bold=True, color=ANSWER_COLOR)
        set_bangla_font(p.add_run(f"{ans_label}. "))
        if ans_tuple:
            render_parts_to_para(p, split_text_and_omml(ans_tuple[0], strip_label_prefix=ans_label))
    # Spacing after each question
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

# =============================================================================
# MAIN LOGIC
# =============================================================================

def extract_mcqs(paragraphs):
    q_re = re.compile(r'^\s*(\d+)\.\s*(.*)')
    ref_re = re.compile(r'\s*(\[.*?\]|\(.*?\))\s*')
    opt_re = re.compile(r'^\s*(ক|খ|গ|ঘ)\.\s*(.*)')
    ans_re = re.compile(r'^\s*উত্তর[:：]\s*(\w)\.?\s*(.*)')
    mcqs = []
    cur = None
    state = "find_question"
    for para in paragraphs:
        lines = para.text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if state == "find_question":
                m = q_re.match(line)
                if m:
                    cur = {'serial': m.group(1), 'question_meta': [(para, m.group(2))], 'reference': '', 'options_meta': {}, 'answer_label': ''}
                    state = "in_question"
            elif state == "in_question":
                original_line = line
                mref = ref_re.search(line)
                if mref:
                    cur['reference'] = clean_text(mref.group(0).strip())
                    line = ref_re.sub('', line, 1).strip()
                mopt = opt_re.match(line)
                if mopt:
                    cur['options_meta'][mopt.group(1)] = (para, mopt.group(2))
                    state = "in_options"
                elif line:
                    cur['question_meta'].append((para, original_line))
            elif state == "in_options":
                mans = ans_re.match(line)
                mopt = opt_re.match(line)
                if mopt:
                    cur['options_meta'][mopt.group(1)] = (para, mopt.group(2))
                elif mans:
                    cur['answer_label'] = mans.group(1)
                    if cur: mcqs.append(cur)
                    cur = None
                    state = "find_question"
                elif cur and cur['options_meta']:
                    last_label = list(cur['options_meta'])[-1]
                    prev_para, prev_text = cur['options_meta'][last_label]
                    cur['options_meta'][last_label] = (prev_para, f"{prev_text} {line}")
    return mcqs

def main():
    src_file = "AP_-_MCQ_Sheet_-_Class_6_-_Chapter_1.1^J_1.2^J_1.3^J_1.4^J_1.5^J_1.6_-_স্বাভাবিক_সংখ্যা_ও_ভগ্নাংশ.docx"
    out_file = "Reformatted_MCQ_Sheet.docx"
    print(f"Starting conversion of '{src_file}'...")
    try:
        doc = Document(src_file)
    except Exception as e:
        print(f"Error: Could not open source file. {e}")
        return
    outdoc = Document()
    # Setup 2-column layout
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    # --- Set tab stops for options alignment: change these to fit your font visually ---
    # Example: each option at every 2 inches (for 4 columns)
    tab_stops_oneline = [Inches(0.8), Inches(1.6), Inches(2.4)]    # for ক খ গ ঘ (edit as you wish)
    tab_stops_twoline = [Inches(1.6)]                              # for ক/খ and গ/ঘ, single tab

    mcqs = extract_mcqs(doc.paragraphs)
    for mcq in mcqs:
        try:
            layout = option_layout_type(mcq['options_meta'])
            if layout == "oneline":
                format_mcq(mcq, outdoc, tab_stops_oneline, layout)
            elif layout == "twoline":
                format_mcq(mcq, outdoc, tab_stops_twoline, layout)
            else:
                format_mcq(mcq, outdoc, [], layout)
        except Exception as e:
            print(f"Error formatting MCQ #{mcq.get('serial', 'N/A')}. Skipping. Error: {e}")

    outdoc.save(out_file)
    print(f"Success! Output saved to '{out_file}'")

if __name__ == "__main__":
    main()
