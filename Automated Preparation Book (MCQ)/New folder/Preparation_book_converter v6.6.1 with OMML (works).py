import re
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches

# --- Configuration Constants ---
FONT_NAME = "Tiro Bangla"
FONT_SIZE = 11
ANSWER_COLOR = RGBColor(0x08, 0x85, 0x65)
OPTION_LABELS = ['ক', 'খ', 'গ', 'ঘ']

# Option length (character) thresholds for layout selection
SHORT_OPTION_CHAR_LIMIT = 5    # 1 line (all options <= this)
MEDIUM_OPTION_CHAR_LIMIT = 15  # 2 lines (all options <= this)
LONG_OPTION_CHAR_LIMIT = 20    # 4 lines (any option > this)
OMML_EXTRA_WEIGHT = 18         # treat each OMML as this many chars (tune for your MCQs)
ROMAN_NUMERALS = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']

def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'([(\[\{«“‘])\s+', r'\1', text)
    text = re.sub(r'\s+([)\]\}»”’.,?।;:])', r'\1', text)
    text = re.sub(r' {2,}', ' ', text)
    return text

def patch_omml_font_size(omml_xml, size_pt=FONT_SIZE):
    try:
        size_val = str(int(size_pt * 2))
        tree = etree.fromstring(omml_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for sz_tag in tree.xpath('.//w:sz|.//w:szCs', namespaces=ns):
            sz_tag.attrib[qn('w:val')] = size_val
        return etree.tostring(tree, encoding='unicode')
    except etree.XMLSyntaxError:
        return omml_xml

def set_bangla_font(run, font_size=FONT_SIZE, bold=False, color=None):
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def split_text_and_omml(para, strip_label_prefix=None):
    para_xml = etree.tostring(para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    parts = []
    first_text_run = True
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            text = "".join(t.text or "" for t in node.findall('.//w:t', namespaces=ns))
            if first_text_run and strip_label_prefix:
                regex = re.compile(r'^\s*' + re.escape(strip_label_prefix) + r'\.\s*')
                if regex.match(text):
                    text = regex.sub('', text, 1)
                    first_text_run = False
            if text:
                parts.append(("text", text))
        elif node.tag.endswith(('oMath', 'oMathPara')):
            omml_xml = etree.tostring(node, encoding='unicode')
            parts.append(("omml", omml_xml))
    return parts

def render_parts_to_para(para, parts, bold=False, color=None):
    for i, (ctype, cvalue) in enumerate(parts):
        if i > 0:
            prev_ctype, _ = parts[i-1]
            if prev_ctype != ctype:
                para.add_run(" ")
        if ctype == "text":
            run = para.add_run(clean_text(cvalue))
            set_bangla_font(run, bold=bold, color=color)
        elif ctype == "omml":
            omml_patched = patch_omml_font_size(cvalue)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_patched +
                '</w:r>'
            )
            para._p.append(omml_run)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)

def get_option_length_class(options_dict):
    """
    Decide layout by all 4 options:
        - 'oneline'  : all options short (no long OMML)
        - 'twoline'  : all <= MEDIUM_OPTION_CHAR_LIMIT (text+OMML), else...
        - 'fourline' : any option long (text or OMML)
    """
    maxlen = 0
    all_short = True
    all_medium = True
    for label in OPTION_LABELS:
        opt_tuple = options_dict.get(label)
        if not opt_tuple:
            return "fourline"
        para, _ = opt_tuple
        option_len = 0
        for ctype, cval in split_text_and_omml(para, strip_label_prefix=label):
            if ctype == "text":
                option_len += len(cval.strip())
            elif ctype == "omml":
                # OMML gets counted as some extra "weight"
                option_len += OMML_EXTRA_WEIGHT
        maxlen = max(maxlen, option_len)
        if option_len > SHORT_OPTION_CHAR_LIMIT:
            all_short = False
        if option_len > MEDIUM_OPTION_CHAR_LIMIT:
            all_medium = False
    if all_short:
        return "oneline"
    elif all_medium:
        return "twoline"
    else:
        return "fourline"

def format_mcq(mcq, doc, tab_stops_oneline, tab_stops_twoline):
    q_para = doc.add_paragraph()
    set_bangla_font(q_para.add_run(f"{mcq['serial']}. "), bold=True)
    prompt_line = None
    for i, (para, text) in enumerate(mcq['question_meta']):
        if i == 0:
            render_parts_to_para(q_para, split_text_and_omml(para, strip_label_prefix=mcq['serial']))
        else:
            if any(key in text for key in ["সঠিক", "ঠিক", "যথাযথ", " কোনটি", "কোনগুলো"]):
                prompt_line = para
            else:
                sub_para = doc.add_paragraph()
                set_bangla_font(sub_para.add_run(f"{ROMAN_NUMERALS[i-1]} "), bold=True)
                render_parts_to_para(sub_para, split_text_and_omml(para))
    if mcq.get('reference'):
        q_para.add_run(" ")
        set_bangla_font(q_para.add_run(mcq['reference']))
    if prompt_line:
        render_parts_to_para(doc.add_paragraph(), split_text_and_omml(prompt_line))
    # --- Option Layout ---
    option_layout = get_option_length_class(mcq['options_meta'])
    if option_layout == "oneline":
        option_para = doc.add_paragraph()
        option_para.paragraph_format.tab_stops.clear_all()
        for tab_pos in tab_stops_oneline:
            option_para.paragraph_format.tab_stops.add_tab_stop(tab_pos)
        col = 0
        for label in OPTION_LABELS:
            opt_tuple = mcq['options_meta'].get(label)
            if not opt_tuple:
                continue
            if col > 0:
                option_para.add_run('\t')
            run = option_para.add_run(f"{label}. ")
            set_bangla_font(run, bold=True)
            render_parts_to_para(option_para, split_text_and_omml(opt_tuple[0], strip_label_prefix=label))
            col += 1
    elif option_layout == "twoline":
        for row in [(0, 1), (2, 3)]:
            option_para = doc.add_paragraph()
            option_para.paragraph_format.tab_stops.clear_all()
            for tab_pos in tab_stops_twoline:
                option_para.paragraph_format.tab_stops.add_tab_stop(tab_pos)
            col = 0
            for idx in row:
                label = OPTION_LABELS[idx]
                opt_tuple = mcq['options_meta'].get(label)
                if not opt_tuple:
                    continue
                if col > 0:
                    option_para.add_run('\t')
                run = option_para.add_run(f"{label}. ")
                set_bangla_font(run, bold=True)
                render_parts_to_para(option_para, split_text_and_omml(opt_tuple[0], strip_label_prefix=label))
                col += 1
    else:
        for idx in range(4):
            label = OPTION_LABELS[idx]
            opt_tuple = mcq['options_meta'].get(label)
            if not opt_tuple:
                continue
            p = doc.add_paragraph()
            run = p.add_run(f"{label}. ")
            set_bangla_font(run, bold=True)
            render_parts_to_para(p, split_text_and_omml(opt_tuple[0], strip_label_prefix=label))
    # Answer Line
    ans_label = mcq.get('answer_label')
    if ans_label:
        ans_tuple = mcq['options_meta'].get(ans_label)
        p = doc.add_paragraph()
        set_bangla_font(p.add_run("উত্তর: "), bold=True, color=ANSWER_COLOR)
        set_bangla_font(p.add_run(f"{ans_label}. "))
        if ans_tuple:
            render_parts_to_para(p, split_text_and_omml(ans_tuple[0], strip_label_prefix=ans_label))
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

def extract_mcqs(paragraphs):
    q_re = re.compile(r'^\s*(\d+)\.\s*(.*)')
    ref_re = re.compile(r'\s*(\[.*?\]|\(.*?\))\s*')
    opt_re = re.compile(r'^\s*(ক|খ|গ|ঘ)\.\s*(.*)')
    ans_re = re.compile(r'^\s*উত্তর[:：]\s*(\w)\.?\s*(.*)')
    mcqs = []
    cur = None
    state = "find_question"
    for para in paragraphs:
        lines = para.text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if state == "find_question":
                m = q_re.match(line)
                if m:
                    cur = {'serial': m.group(1), 'question_meta': [(para, m.group(2))], 'reference': '', 'options_meta': {}, 'answer_label': ''}
                    state = "in_question"
            elif state == "in_question":
                original_line = line
                mref = ref_re.search(line)
                if mref:
                    cur['reference'] = clean_text(mref.group(0).strip())
                    line = ref_re.sub('', line, 1).strip()
                mopt = opt_re.match(line)
                if mopt:
                    cur['options_meta'][mopt.group(1)] = (para, mopt.group(2))
                    state = "in_options"
                elif line:
                    cur['question_meta'].append((para, original_line))
            elif state == "in_options":
                mans = ans_re.match(line)
                mopt = opt_re.match(line)
                if mopt:
                    cur['options_meta'][mopt.group(1)] = (para, mopt.group(2))
                elif mans:
                    cur['answer_label'] = mans.group(1)
                    if cur: mcqs.append(cur)
                    cur = None
                    state = "find_question"
                elif cur and cur['options_meta']:
                    last_label = list(cur['options_meta'])[-1]
                    prev_para, prev_text = cur['options_meta'][last_label]
                    cur['options_meta'][last_label] = (prev_para, f"{prev_text} {line}")
    return mcqs

def main():
    src_file = "AP_-_MCQ_Sheet_-_Class_6_-_Chapter_1.1^J_1.2^J_1.3^J_1.4^J_1.5^J_1.6_-_স্বাভাবিক_সংখ্যা_ও_ভগ্নাংশ.docx"
    out_file = "Reformatted_MCQ_Sheet.docx"
    print(f"Starting conversion of '{src_file}'...")
    try:
        doc = Document(src_file)
    except Exception as e:
        print(f"Error: Could not open source file. {e}")
        return
    outdoc = Document()
    # Setup 2-column layout
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    # Tab stops for your desired alignment
    tab_stops_oneline = [Inches(0.8), Inches(1.6), Inches(2.4)]
    tab_stops_twoline = [Inches(1.6)]

    mcqs = extract_mcqs(doc.paragraphs)
    print(f"Found and parsed {len(mcqs)} MCQs.")

    for mcq in mcqs:
        try:
            format_mcq(mcq, outdoc, tab_stops_oneline, tab_stops_twoline)
        except Exception as e:
            print(f"Error formatting MCQ #{mcq.get('serial', 'N/A')}. Skipping. Error: {e}")

    outdoc.save(out_file)
    print(f"Success! Output saved to '{out_file}'")

if __name__ == "__main__":
    main()
