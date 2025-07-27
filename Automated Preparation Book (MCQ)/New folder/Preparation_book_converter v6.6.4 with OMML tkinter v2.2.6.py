import re
import tkinter as tk
from tkinter import filedialog, messagebox
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches
import pandas as pd
import json

# --- Constants and Configuration ---
FONT_NAME = "Tiro Bangla"
FONT_SIZE = 11
ANSWER_COLOR = RGBColor(0x08, 0x85, 0x65)
OPTION_LABELS_BN = ['ক', 'খ', 'গ', 'ঘ']
OPTION_LABELS_EN = ['a', 'b', 'c', 'd', 'A', 'B', 'C', 'D']
SHORT_OPTION_CHAR_LIMIT = 4
MEDIUM_OPTION_CHAR_LIMIT = 18
LONG_OPTION_CHAR_LIMIT = 1000
OMML_WEIGHT_SHORT = 4
OMML_WEIGHT_MEDIUM = 12
OMML_WEIGHT_LONG = 1000
ROMAN_NUMERALS = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']

# --- Core Functions ---

def clean_text(text):
    """Cleans up whitespace and punctuation in a combined string."""
    if not text:
        return ""
    # Normalize all whitespace (including newlines, tabs) to a single space
    text = re.sub(r'\s+', ' ', text)
    # Fix spacing around punctuation
    text = re.sub(r'\s*([)\]\}»”’.,?।;:])\s*', r'\1', text)
    text = re.sub(r'\s*([(\[\{«“‘])\s*', r'\1', text)
    return text.strip()

def patch_omml_font_size(omml_xml, size_pt=FONT_SIZE):
    """Sets a consistent font size for OMML (equation) elements."""
    try:
        size_val = str(int(size_pt * 2))
        tree = etree.fromstring(omml_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for sz_tag in tree.xpath('.//w:sz|.//w:szCs', namespaces=ns):
            sz_tag.attrib[qn('w:val')] = size_val
        return etree.tostring(tree, encoding='unicode')
    except etree.XMLSyntaxError:
        return omml_xml

def set_bangla_font(run, font_size=FONT_SIZE, bold=False, color=None):
    """Applies standard Bangla font styling to a run."""
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
# --- NEW: Advanced LaTeX Equation Converter ---
MML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
LATEX_SYMBOLS = {'…': r'\dots', '≠': r'\neq', '≤': r'\leq', '≥': r'\geq', '×': r'\times', '÷': r'\div', '±': r'\pm', '∞': r'\infty', '°': r'^{\circ}', 'α': r'\alpha', 'β': r'\beta', 'γ': r'\gamma', 'δ': r'\delta', 'ε': r'\epsilon', 'θ': r'\theta', 'λ': r'\lambda', 'μ': r'\mu', 'π': r'\pi', 'ρ': r'\rho', 'σ': r'\sigma', 'φ': r'\phi', 'ω': r'\omega', 'Δ': r'\Delta', 'Ω': r'\Omega', '∑': r'\sum', '∫': r'\int'}
KNOWN_FUNCTIONS = ['sin', 'cos', 'tan', 'log', 'ln', 'lim', 'exp']

def _parse_omml_to_latex_recursive(node):
    if node is None: return ""
    tag = etree.QName(node.tag).localname
    if tag == 't':
        text = node.text or ''
        for char, command in LATEX_SYMBOLS.items():
            text = text.replace(char, f" {command} ")
        return f"\\{text.strip()}" if text.strip() in KNOWN_FUNCTIONS else text
    elif tag == 'f':
        num = _parse_omml_to_latex_recursive(node.find(f'./{{{MML_NS}}}num'))
        den = _parse_omml_to_latex_recursive(node.find(f'./{{{MML_NS}}}den'))
        return f"\\frac{{{num.strip()}}}{{{den.strip()}}}"
    elif tag == 'sSup':
        base = _parse_omml_to_latex_recursive(node.find(f'./{{{MML_NS}}}e'))
        sup = _parse_omml_to_latex_recursive(node.find(f'./{{{MML_NS}}}sup'))
        return f"{{{base.strip()}}}^{{{sup.strip()}}}"
    elif tag == 'sSub':
        base = _parse_omml_to_latex_recursive(node.find(f'./{{{MML_NS}}}e'))
        sub = _parse_omml_to_latex_recursive(node.find(f'./{{{MML_NS}}}sub'))
        return f"{{{base.strip()}}}_{{{sub.strip()}}}"
    elif tag == 'rad':
        base = _parse_omml_to_latex_recursive(node.find(f'./{{{MML_NS}}}e'))
        return f"\\sqrt{{{base.strip()}}}"
    else:
        return "".join(_parse_omml_to_latex_recursive(child) for child in node)

def omml_to_latex(omml_xml):
    try:
        tree = etree.fromstring(omml_xml)
        math_node = tree.find(f'.//{{{MML_NS}}}oMath') or tree.find(f'.//{{{MML_NS}}}oMathPara')
        if math_node is None:
            if etree.QName(tree.tag).localname in ('oMath', 'oMathPara'): math_node = tree
            else: return ""
        latex_str = _parse_omml_to_latex_recursive(math_node).strip()
        latex_str = re.sub(r'\s+', ' ', latex_str)
        return f"${latex_str}$"
    except Exception:
        return ""

def get_para_content(para, use_latex=False):
    """Gets paragraph content, optionally converting OMML to LaTeX."""
    full_str = ""
    converter_func = omml_to_latex if use_latex else (lambda xml: etree.fromstring(xml).xpath('string()'))
    for ctype, cvalue in split_text_and_omml(para):
        if ctype == "text":
            full_str += cvalue
        elif ctype == "omml":
            full_str += converter_func(cvalue)
    return full_str.strip()
    
# --- Parsing and Extraction (Logic for Type and Reference added) ---
def split_text_and_omml(para, strip_label_prefix=None):
    para_xml = etree.tostring(para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    parts, chars_to_skip = [], 0
    full_text = "".join(tree.xpath('.//w:t/text()', namespaces=ns))
    if strip_label_prefix:
        match = re.match(r'^\s*[\(\[]?\s*%s\s*[\.\)\]\।:ঃ]?\s*' % re.escape(strip_label_prefix), full_text, re.IGNORECASE)
        if match: chars_to_skip = len(match.group(0))
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            for t in node.findall('.//w:t', namespaces=ns):
                if chars_to_skip > 0 and t.text:
                    if len(t.text) <= chars_to_skip: chars_to_skip -= len(t.text); t.text = ""
                    else: t.text = t.text[chars_to_skip:]; chars_to_skip = 0
            final_run_text = "".join(t.text or "" for t in node.findall('.//w:t', namespaces=ns))
            if final_run_text: parts.append(("text", final_run_text))
        elif node.tag.endswith(('oMath', 'oMathPara')):
            if chars_to_skip > 0: chars_to_skip = 0
            parts.append(("omml", etree.tostring(node, encoding='unicode')))
    return parts


def render_parts_to_para(para, parts, bold=False, color=None):
    for i, (ctype, cvalue) in enumerate(parts):
        # Add a space between text and omml when needed
        if i > 0:
            prev_ctype, prev_cvalue = parts[i-1]
            # If previous is text and does not end with space and current is equation, or
            # previous is equation and current is text and current does not start with space
            need_space = (
                (prev_ctype == "text" and ctype == "omml" and not prev_cvalue.endswith(" ")) or
                (prev_ctype == "omml" and ctype == "text" and not cvalue.startswith(" "))
            )
            if need_space:
                para.add_run(" ")
        if ctype == "text":
            run = para.add_run(cvalue)
            set_bangla_font(run, bold=bold, color=color)
        elif ctype == "omml":
            omml_patched = patch_omml_font_size(cvalue)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_patched +
                '</w:r>'
            )
            para._p.append(omml_run)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)


# --- Length Calculation and Parsing Functions ---

def omml_to_text(omml_xml):
    try:
        tree = etree.fromstring(omml_xml)
        # get all math 't' (text) nodes
        math_texts = tree.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}t')
        text = ''.join([t.text or '' for t in math_texts])
        return text
    except Exception:
        return ""

def _omml_visible_length(omml_xml):
    try:
        tree = etree.fromstring(omml_xml)
        math_texts = tree.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}t')
        text = ''.join([t.text or '' for t in math_texts])
        return len(text)
    except Exception:
        return 12

def _option_effective_length(para, label):
    total = 0
    for ctype, cval in split_text_and_omml(para, strip_label_prefix=label):
        if ctype == "text":
            total += len(cval.strip())
        elif ctype == "omml":
            visible_len = _omml_visible_length(cval)
            if visible_len <= SHORT_OPTION_CHAR_LIMIT:
                total += OMML_WEIGHT_SHORT
            elif visible_len <= MEDIUM_OPTION_CHAR_LIMIT:
                total += OMML_WEIGHT_MEDIUM
            else:
                total += OMML_WEIGHT_LONG
    return total

def get_option_length_class(options_dict):
    if not options_dict: return "fourline"
    lens = [_option_effective_length(p, l) for l, (p, _) in options_dict.items() if p]
    if not lens: return "fourline"
    if all(l <= SHORT_OPTION_CHAR_LIMIT for l in lens): return "oneline"
    if all(l <= MEDIUM_OPTION_CHAR_LIMIT for l in lens): return "twoline"
    return "fourline"

def get_para_full_text(para):
    """
    Extracts a complete text string from a paragraph, including from OMML.
    Used for reliable parsing.
    """
    full_str = ""
    # Use split_text_and_omml without stripping anything to get all content
    parts = split_text_and_omml(para, strip_label_prefix=None)
    for ctype, cvalue in parts:
        if ctype == "text":
            full_str += cvalue
        elif ctype == "omml":
            # Convert OMML to its plain text representation
            full_str += omml_to_text(cvalue)
    return full_str.strip()


def parse_serial_and_question(line):
    m = re.match(r'^\s*[\(]?([০-৯0-9]+)[\.\)\।]?\s*(.*)', line)
    return m.groups() if m else (None, line)

def parse_option(line):
    m = re.match(r'^\s*[\(\[]?\s*([কখগঘa-dA-D])[\.\)\]\।]?\s*(.*)', line)
    return (m.group(1), m.group(2).strip()) if m else (None, line)

def parse_answer(line):
    m = re.match(r'^উত্তর[:：ঃ]?\s*[\(\[]?\s*([কখগঘa-dA-D])[\.\)\]\।]?\s*(.*)', line, re.IGNORECASE)
    if m:
        return (m.group(1), m.group(2).lstrip(').।. ').strip())
    m2 = re.match(r'^উত্তর[:：ঃ]?\s*(.*)', line, re.IGNORECASE)
    return (None, m2.group(1).lstrip(').।. ').strip()) if m2 else (None, None)

def parse_explanation(line):
    m = re.match(r'^(ব্যাখ্যা[:：ঃ]?)\s*(.*)', line)
    return (m.group(1), m.group(2)) if m else (None, None)


def extract_mcqs(paragraphs):
    """Extracts all structured MCQ data from a list of paragraphs."""
    mcqs, para_to_mcq, cur, state = [], {}, None, "find_question"
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        line = get_para_content(para)
        if not line and state != "in_options":
            i += 1
            continue

        if state == "find_question":
            serial_match = re.match(r'^\s*[\(]?([০-৯0-9]+)[\.\)\।]?', line)
            if serial_match:
                serial = serial_match.group(1)
                cur = {'serial': serial, 'question_meta': [], 'reference': '', 'qtype': 'সাধারণ', 'options_meta': {}, 'answer_label': None, 'answer_meta': None, 'explanation_meta': None, 'all_paras': {i}}
                para_to_mcq[i] = len(mcqs)
                state = "in_question"
                # This paragraph is part of the question, so don't increment i yet
                continue 
            i += 1
        
        elif state == "in_question":
            is_option = re.match(r'^\s*[\(\[]?\s*([কখগঘa-dA-D])[\.\)\]\।]?', line)
            if is_option:
                state = "in_options"
                continue

            # Capture all references like [tag1] [tag2]
            mrefs = re.findall(r'(\[.*?\])', para.text)
            if mrefs:
                cur['reference'] += " ".join(mrefs)

            cur['question_meta'].append((para, None))
            cur['all_paras'].add(i)
            para_to_mcq[i] = len(mcqs)
            i += 1

        elif state == "in_options":
            opt_match = re.match(r'^\s*[\(\[]?\s*([কখগঘa-dA-D])[\.\)\]\।]?\s*(.*)', line)
            ans_match = re.match(r'^উত্তর[:：ঃ]?\s*[\(\[]?\s*([কখগঘa-dA-D])[\.\)\]\।]?\s*(.*)', line, re.IGNORECASE)
            exp_match = re.match(r'^(ব্যাখ্যা[:：ঃ]?)\s*(.*)', line)

            if opt_match:
                cur['options_meta'][opt_match.group(1)] = (para, opt_match.group(2).strip())
            elif ans_match:
                cur['answer_label'] = ans_match.group(1)
                cur['answer_meta'] = (para, ans_match.group(2).lstrip(').। ').strip())
            elif exp_match:
                cur['explanation_meta'] = (para, exp_match.group(1))
            
            # This is the end of an MCQ block
            if ans_match or exp_match:
                # --- Re-evaluate Question Type before saving ---
                is_bahupadi = any(any(p.text.strip().startswith(r) for r in ROMAN_NUMERALS) for p, _ in cur['question_meta'])
                has_stem = len(cur['question_meta']) > 1 or 'উদ্দীপক' in get_para_content(cur['question_meta'][0][0])
                if is_bahupadi:
                    cur['qtype'] = 'বহুপদী সমাপ্তিসূচক'
                elif has_stem:
                    cur['qtype'] = 'উদ্দীপকভিত্তিক'
                
                mcqs.append(cur)
                state = "find_question"
                cur = None
            
            i += 1
            
    if cur: mcqs.append(cur) # Save the last one if file ends unexpectedly
    return mcqs, para_to_mcq

# --- Document Formatting and Generation ---

def get_full_content_string(para, label_to_strip):
    """
    Extracts a normalized string from a paragraph, including text from OMML,
    after stripping a given prefix/label. Used for content comparison.
    """
    full_str = ""
    parts = split_text_and_omml(para, strip_label_prefix=label_to_strip)
    for ctype, cvalue in parts:
        if ctype == "text":
            full_str += cvalue
        elif ctype == "omml":
            full_str += omml_to_text(cvalue)
    # Normalize by removing all whitespace and making it lowercase for robust comparison
    return re.sub(r'\s+', '', full_str).lower()

def format_mcq(mcq, doc, tab_stops_oneline, tab_stops_twoline):
    """Formats and writes a single MCQ to the output document."""
    
    # --- Question ---
    q_para = doc.add_paragraph()
    set_bangla_font(q_para.add_run(f"{mcq['serial']}. "), bold=True)
    
    # --- Question and Metadata Processing ---
    full_question_parts = []
    prompt_line_para = None
    roman_numeral_paras = []
    
    # Process the main question paragraph first
    main_question_para, _ = mcq['question_meta'][0]
    full_question_parts.extend(split_text_and_omml(main_question_para, strip_label_prefix=mcq['serial']))

    # Process subsequent paragraphs associated with the question
    for p, _ in mcq['question_meta'][1:]:
        text = (p.text or '').strip()
        # Skip pure reference lines, as they are handled separately
        if text.startswith("[") and text.endswith("]") and text.count("[") == 1:
            continue
        elif any(key in text for key in ["সঠিক", "ঠিক", "যথাযথ", " কোনটি", "কোনগুলো"]):
            prompt_line_para = p
        else:
            roman_numeral_paras.append(p)

    render_parts_to_para(q_para, full_question_parts)
    
    # Add all collected reference tags at the end of the question line
    ref = mcq.get('reference', '').strip()
    if ref:
        q_para.add_run("  ") # Extra space for clarity
        set_bangla_font(q_para.add_run(ref))

    # Render roman numeral items if any exist
    for i, p in enumerate(roman_numeral_paras):
        sub_p = doc.add_paragraph()
        set_bangla_font(sub_p.add_run(f"{ROMAN_NUMERALS[i]} "), bold=True)
        render_parts_to_para(sub_p, split_text_and_omml(p))

    # Render the prompt line (e.g., "Which is correct?")
    if prompt_line_para:
        render_parts_to_para(doc.add_paragraph(), split_text_and_omml(prompt_line_para))

    # --- Options ---
    option_layout = get_option_length_class(mcq['options_meta'])
    # Ensure a consistent order for options
    label_order = [lbl for lbl in OPTION_LABELS_BN if lbl in mcq['options_meta']] + \
                  [lbl for lbl in OPTION_LABELS_EN if lbl in mcq['options_meta']]
    label_map_bn = dict(zip(label_order, OPTION_LABELS_BN))

    def render_option(p, label, bn_label):
        opt_tuple = mcq['options_meta'].get(label)
        if not opt_tuple: return
        set_bangla_font(p.add_run(f"{bn_label}. "), bold=True)
        parts = split_text_and_omml(opt_tuple[0], strip_label_prefix=label)
        render_parts_to_para(p, parts)
        
    if option_layout == "oneline":
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.clear_all()
        for tab in tab_stops_oneline: p.paragraph_format.tab_stops.add_tab_stop(tab)
        for i, label in enumerate(label_order[:4]):
            if i > 0: p.add_run('\t')
            render_option(p, label, label_map_bn.get(label))
            
    elif option_layout == "twoline":
        for i, row in enumerate([(0, 1), (2, 3)]):
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.clear_all()
            for tab in tab_stops_twoline: p.paragraph_format.tab_stops.add_tab_stop(tab)
            for j, opt_idx in enumerate(row):
                if opt_idx >= len(label_order): continue
                if j > 0: p.add_run('\t')
                label = label_order[opt_idx]
                render_option(p, label, label_map_bn.get(label))
    else: # "fourline"
        for i, label in enumerate(label_order[:4]):
             p = doc.add_paragraph()
             render_option(p, label, label_map_bn.get(label))

    # --- Answer Section (REVISED LOGIC) ---
    ans_label = mcq.get('answer_label')
    ans_meta = mcq.get('answer_meta') # Tuple: (paragraph, raw_text)

    if ans_label or ans_meta:
        p = doc.add_paragraph()
        set_bangla_font(p.add_run("উত্তর: "), bold=True, color=ANSWER_COLOR)

        # Case 1: The answer is given with an explicit label (e.g., "উত্তর: ক")
        if ans_label:
            bn_label = label_map_bn.get(ans_label, ans_label)
            set_bangla_font(p.add_run(f"{bn_label}. "))
            ans_tuple = mcq['options_meta'].get(ans_label)
            if ans_tuple:
                parts = split_text_and_omml(ans_tuple[0], strip_label_prefix=ans_label)
                render_parts_to_para(p, parts)
            elif ans_meta: # Fallback
                _, ans_text_from_meta = ans_meta
                set_bangla_font(p.add_run(ans_text_from_meta))

        # Case 2: No label given. We must find which option matches the content.
        elif ans_meta:
            ans_para, _ = ans_meta
            answer_content_str = get_full_content_string(ans_para, "উত্তর")
            
            found_label = None
            for label, (opt_para, _) in mcq['options_meta'].items():
                option_content_str = get_full_content_string(opt_para, label)
                
                if answer_content_str and answer_content_str == option_content_str:
                    found_label = label
                    break
            
            if found_label:
                bn_label = label_map_bn.get(found_label, found_label)
                set_bangla_font(p.add_run(f"{bn_label}. "))
                matched_para, _ = mcq['options_meta'][found_label]
                parts = split_text_and_omml(matched_para, strip_label_prefix=found_label)
                render_parts_to_para(p, parts)
            else:
                # Fallback: couldn't match. Render original answer line content.
                parts = split_text_and_omml(ans_para, strip_label_prefix="উত্তর")
                render_parts_to_para(p, parts)
                
    # --- Explanation Section ---
    if mcq.get('explanation_meta'):
        para, label = mcq['explanation_meta']
        ep = doc.add_paragraph()
        set_bangla_font(ep.add_run("ব্যাখ্যা: "), bold=True)
        parts = split_text_and_omml(para, strip_label_prefix="ব্যাখ্যা")
        render_parts_to_para(ep, parts)

    # Add space after the entire MCQ block
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(8)


def copy_para_with_omml(src_para, outdoc, bold=False):
    out_p = outdoc.add_paragraph()
    render_parts_to_para(out_p, split_text_and_omml(src_para), bold=bold)
    return out_p

def is_extra_heading(text):
    keywords = ["উদ্দীপক", "নিচের", "তথ্যের ভিত্তিতে", "তথ্যের আলোকে", "প্রশ্নের উত্তর দাও"]
    return any(text.strip().startswith(kw) for kw in keywords)

# MML namespace constant for cleaner code
MML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def to_unicode_math(text):
    """Converts standard latin characters to their mathematical italic unicode counterparts."""
    def _math_italic(c):
        if 'a' <= c <= 'z':
            return chr(ord('\U0001D44E') + ord(c) - ord('a'))
        if 'A' <= c <= 'Z':
            return chr(ord('\U0001D434') + ord(c) - ord('A'))
        return c
    return ''.join(_math_italic(c) for c in text)

def _parse_omml_node_recursive(node):
    """
    Recursively traverses an OMML XML node to build a properly formatted string.
    This is the core of the fix.
    """
    if node is None:
        return ""

    tag = etree.QName(node.tag).localname

    # Base case: A text node
    if tag == 't':
        # Convert text to mathematical italics and return
        return to_unicode_math(node.text or '')

    # Structural nodes that require specific formatting
    elif tag == 'f':  # Fraction
        num = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}num'))
        den = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}den'))
        return f"({num})/({den})"

    elif tag == 'sSup':  # Superscript (e.g., x^2)
        base = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}e'))
        sup = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}sup'))
        return f"{base}^{{{sup}}}"

    elif tag == 'sSub':  # Subscript (e.g., x_1)
        base = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}e'))
        sub = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}sub'))
        return f"{base}_{{{sub}}}"

    elif tag == 'rad':  # Radical (Square Root)
        base = _parse_omml_node_recursive(node.find(f'./{{{MML_NS}}}e'))
        return f"sqrt({base})"

    # For all other container nodes (like 'r', 'e', 'oMath'),
    # simply process their children and join the results.
    else:
        return "".join(_parse_omml_node_recursive(child) for child in node)

def omml_to_latex_text(omml_xml):
    """
    Converts OMML XML to a clean, human-readable plain text string for export.
    This function replaces the original problematic one.
    """
    try:
        tree = etree.fromstring(omml_xml)
        # Find the main math element, which is usually <m:oMath>
        math_node = tree.find(f'.//{{{MML_NS}}}oMath')
        if math_node is None:
            math_node = tree.find(f'.//{{{MML_NS}}}oMathPara')

        # Handle cases where the root element itself is the math node
        if math_node is None and etree.QName(tree.tag).localname in ('oMath', 'oMathPara'):
            math_node = tree

        if math_node is not None:
            # Start the recursive parsing from the main math node
            return _parse_omml_node_recursive(math_node)
        else:
            # Fallback to simple text extraction if parsing fails
            return omml_to_text(omml_xml)
            
    except Exception:
        # Fallback for any unexpected error
        return omml_to_text(omml_xml)

def mcqs_to_rows(mcqs):
    """Converts MCQ list to rows, now with Type and Reference columns."""
    rows = []
    for mcq in mcqs:
        # Build the full question text from all its paragraphs
        question_text = " ".join(get_para_content(p, use_latex=True) for p, _ in mcq['question_meta'])
        
        # Get the reference text
        reference_text = mcq.get('reference', '').strip()
        
        # Clean the reference from the question text
        if reference_text:
            ref_tags = re.findall(r'(\[.*?\])', reference_text)
            for tag in ref_tags:
                question_text = question_text.replace(tag, '')
        
        # Clean the serial number from the start
        question = re.sub(r'^\s*[\(]?\s*[\d০-৯]+[\.\)\।]?\s*', '', question_text).strip()

        options = {}
        for label in OPTION_LABELS_BN:
            opt_tuple = mcq['options_meta'].get(label)
            if opt_tuple:
                opt_text = get_para_content(opt_tuple[0], use_latex=True)
                options[label] = re.sub(fr'^\s*[\(\[]?{label}[\.\)\]]?\s*', '', opt_text).strip()
            else:
                options[label] = ""

        row = {
            'Serial': mcq['serial'],
            'Type': mcq.get('qtype', 'সাধারণ'),
            'Reference': reference_text,
            'Question': question,
            'Option_ক': options.get('ক', ''),
            'Option_খ': options.get('খ', ''),
            'Option_গ': options.get('গ', ''),
            'Option_ঘ': options.get('ঘ', ''),
            'Answer_Label': mcq.get('answer_label', '')
        }
        rows.append(row)
    return rows

def export_mcqs_to_csv(mcqs, csv_path="mcqs.csv"):
    rows = mcqs_to_rows(mcqs)
    df = pd.DataFrame(rows)
    df.to_csv(csv_path, index=False, encoding='utf-8-sig')

def export_mcqs_to_excel(mcqs, excel_path="mcqs.xlsx"):
    rows = mcqs_to_rows(mcqs)
    df = pd.DataFrame(rows)
    df.to_excel(excel_path, index=False)

def export_mcqs_to_json(mcqs, json_path="mcqs.json"):
    rows = mcqs_to_rows(mcqs)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(rows, f, ensure_ascii=False, indent=2)


def convert_file(src_file, out_file):
    try:
        doc = Document(src_file)
    except Exception as e:
        messagebox.showerror("File Error", f"Could not open source file.\n\n{e}")
        return False
        
    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '210')

    tab_stops_oneline = [Inches(0.8), Inches(1.6), Inches(2.4)]
    tab_stops_twoline = [Inches(1.6)]

    mcqs, para_to_mcq = extract_mcqs(doc.paragraphs)
    export_mcqs_to_csv(mcqs, "your_mcqs.csv")
    export_mcqs_to_excel(mcqs, "your_mcqs.xlsx")
    export_mcqs_to_json(mcqs, "your_mcqs.json")
    processed_mcq_indices = set()
    
    i = 0
    while i < len(doc.paragraphs):
        if i in para_to_mcq:
            mcq_idx = para_to_mcq[i]
            if mcq_idx not in processed_mcq_indices:
                format_mcq(mcqs[mcq_idx], outdoc, tab_stops_oneline, tab_stops_twoline)
                processed_mcq_indices.add(mcq_idx)
            i = max(mcqs[mcq_idx]['all_paras']) + 1
        else:
            para = doc.paragraphs[i]
            # Process non-MCQ paragraphs (headings, etc.)
            if get_para_full_text(para):
                copy_para_with_omml(para, outdoc, bold=is_extra_heading(para.text))
            i += 1
            
    try:
        outdoc.save(out_file)
    except Exception as e:
        messagebox.showerror("Save Error", f"Could not save output file.\n\n{e}")
        return False

    return True


# --- Tkinter GUI (Unchanged) ---
def main_gui():
    def select_input():
        filename = filedialog.askopenfilename(title="Select DOCX file", filetypes=[("Word Files", "*.docx")])
        if filename:
            input_entry.delete(0, tk.END)
            input_entry.insert(0, filename)

    def select_output():
        filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")], title="Save Reformatted DOCX as")
        if filename:
            output_entry.delete(0, tk.END)
            output_entry.insert(0, filename)

    def run_conversion():
        src_file = input_entry.get().strip()
        out_file = output_entry.get().strip()
        if not src_file or not out_file:
            messagebox.showwarning("Input Needed", "Please select both input and output files!")
            return

        export_excel = var_excel.get()
        export_csv = var_csv.get()
        export_json = var_json.get()

        if not (export_excel or export_csv or export_json):
            messagebox.showwarning("Export Needed", "Please select at least one export format (Excel, CSV, or JSON)!")
            return
        
        root.config(cursor="wait")
        root.update()
        try:
            # DOCX convert as before
            ok = convert_file(src_file, out_file)
            # Now extract MCQs and export
            doc = Document(src_file)
            mcqs, para_to_mcq = extract_mcqs(doc.paragraphs)
            basepath = out_file.rsplit('.', 1)[0]  # Use output filename as base for other files

            if export_excel:
                export_mcqs_to_excel(mcqs, basepath + ".xlsx")
            if export_csv:
                export_mcqs_to_csv(mcqs, basepath + ".csv")
            if export_json:
                export_mcqs_to_json(mcqs, basepath + ".json")
        finally:
            root.config(cursor="")
        
        if ok:
            messagebox.showinfo("Done!", f"Successfully converted and exported to selected formats!\nSaved as base: {basepath}")

    root = tk.Tk()
    root.title("MCQ Sheet Formatter (.docx)")

    tk.Label(root, text="Input DOCX:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
    input_entry = tk.Entry(root, width=50)
    input_entry.grid(row=0, column=1, padx=5)
    tk.Button(root, text="Browse...", command=select_input).grid(row=0, column=2, padx=5)

    tk.Label(root, text="Output DOCX:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
    output_entry = tk.Entry(root, width=50)
    output_entry.grid(row=1, column=1, padx=5)
    tk.Button(root, text="Browse...", command=select_output).grid(row=1, column=2, padx=5)

    # --- Export format checkboxes ---
    var_excel = tk.IntVar(value=1)
    var_csv = tk.IntVar(value=1)
    var_json = tk.IntVar(value=0)
    tk.Label(root, text="Export As:").grid(row=2, column=0, sticky="e", padx=5, pady=10)
    tk.Checkbutton(root, text="Excel (.xlsx)", variable=var_excel).grid(row=2, column=1, sticky="w")
    tk.Checkbutton(root, text="CSV (.csv)", variable=var_csv).grid(row=2, column=1, padx=110, sticky="w")
    tk.Checkbutton(root, text="JSON (.json)", variable=var_json).grid(row=2, column=1, padx=210, sticky="w")

    tk.Button(root, text="Convert File", command=run_conversion, bg="#13825c", fg="white", width=16, height=2).grid(row=3, column=1, pady=20)

    root.mainloop()

if __name__ == "__main__":
    main_gui()