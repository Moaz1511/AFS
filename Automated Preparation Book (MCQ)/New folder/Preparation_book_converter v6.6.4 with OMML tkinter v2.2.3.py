import re
import tkinter as tk
from tkinter import filedialog, messagebox
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches

# --- Constants and Configuration ---
FONT_NAME = "Tiro Bangla"
FONT_SIZE = 11
ANSWER_COLOR = RGBColor(0x08, 0x85, 0x65)
OPTION_LABELS_BN = ['ক', 'খ', 'গ', 'ঘ']
OPTION_LABELS_EN = ['a', 'b', 'c', 'd', 'A', 'B', 'C', 'D']
SHORT_OPTION_CHAR_LIMIT = 5
MEDIUM_OPTION_CHAR_LIMIT = 15
LONG_OPTION_CHAR_LIMIT = 20
OMML_WEIGHT_SHORT = 5
OMML_WEIGHT_MEDIUM = 12
OMML_WEIGHT_LONG = 1000
ROMAN_NUMERALS = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']

# --- Core Functions ---

def clean_text(text):
    """Cleans up whitespace and punctuation in a combined string."""
    if not text:
        return ""
    # Normalize all whitespace (including newlines, tabs) to a single space
    text = re.sub(r'\s+', ' ', text)
    # Fix spacing around punctuation
    text = re.sub(r'\s*([)\]\}»”’.,?।;:])\s*', r'\1', text)
    text = re.sub(r'\s*([(\[\{«“‘])\s*', r'\1', text)
    return text.strip()

def patch_omml_font_size(omml_xml, size_pt=FONT_SIZE):
    """Sets a consistent font size for OMML (equation) elements."""
    try:
        size_val = str(int(size_pt * 2))
        tree = etree.fromstring(omml_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for sz_tag in tree.xpath('.//w:sz|.//w:szCs', namespaces=ns):
            sz_tag.attrib[qn('w:val')] = size_val
        return etree.tostring(tree, encoding='unicode')
    except etree.XMLSyntaxError:
        return omml_xml

def set_bangla_font(run, font_size=FONT_SIZE, bold=False, color=None):
    """Applies standard Bangla font styling to a run."""
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def split_text_and_omml(para, strip_label_prefix=None):
    """
    Splits a paragraph into a list of text and OMML parts.
    Robustly strips a prefix, even if it's split across multiple text runs.
    """
    para_xml = etree.tostring(para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    parts = []

    # Get the full text to reliably find the prefix
    full_text = "".join(tree.xpath('.//w:t/text()', namespaces=ns))
    
    chars_to_skip = 0
    if strip_label_prefix:
        regex = re.compile(r'^\s*[\(\[]?\s*%s\s*[\.\)\]\।:ঃ]?\s*' % re.escape(strip_label_prefix), re.IGNORECASE)
        match = regex.match(full_text)
        if match:
            # Calculate how many characters of the raw text to skip
            chars_to_skip = len(match.group(0))

    for node in tree.iterchildren():
        if node.tag.endswith('r'):  # A text run
            run_text_nodes = node.findall('.//w:t', namespaces=ns)
            for t in run_text_nodes:
                if chars_to_skip > 0 and t.text:
                    if len(t.text) <= chars_to_skip:
                        chars_to_skip -= len(t.text)
                        t.text = ""
                    else:
                        t.text = t.text[chars_to_skip:]
                        chars_to_skip = 0
            
            final_run_text = "".join(t.text or "" for t in run_text_nodes)
            if final_run_text:
                parts.append(("text", final_run_text))
        
        elif node.tag.endswith(('oMath', 'oMathPara')):  # An equation
            if chars_to_skip > 0:
                # This indicates an equation is part of a label, which is unlikely.
                # We'll stop skipping characters to avoid deleting an equation.
                chars_to_skip = 0
            omml_xml = etree.tostring(node, encoding='unicode')
            parts.append(("omml", omml_xml))

    return parts


def render_parts_to_para(para, parts, bold=False, color=None):
    for i, (ctype, cvalue) in enumerate(parts):
        # Add a space between text and omml when needed
        if i > 0:
            prev_ctype, prev_cvalue = parts[i-1]
            # If previous is text and does not end with space and current is equation, or
            # previous is equation and current is text and current does not start with space
            need_space = (
                (prev_ctype == "text" and ctype == "omml" and not prev_cvalue.endswith(" ")) or
                (prev_ctype == "omml" and ctype == "text" and not cvalue.startswith(" "))
            )
            if need_space:
                para.add_run(" ")
        if ctype == "text":
            # Do NOT run clean_text here; preserve original
            run = para.add_run(cvalue)
            set_bangla_font(run, bold=bold, color=color)
        elif ctype == "omml":
            omml_patched = patch_omml_font_size(cvalue)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_patched +
                '</w:r>'
            )
            para._p.append(omml_run)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)


# --- Length Calculation and Parsing Functions (mostly unchanged) ---

def _omml_visible_length(omml_xml):
    try:
        tree = etree.fromstring(omml_xml)
        math_texts = tree.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}t')
        text = ''.join([t.text or '' for t in math_texts])
        return len(text)
    except Exception:
        return 12

def _option_effective_length(para, label):
    total = 0
    for ctype, cval in split_text_and_omml(para, strip_label_prefix=label):
        if ctype == "text":
            total += len(cval.strip())
        elif ctype == "omml":
            visible_len = _omml_visible_length(cval)
            if visible_len <= SHORT_OPTION_CHAR_LIMIT:
                total += OMML_WEIGHT_SHORT
            elif visible_len <= MEDIUM_OPTION_CHAR_LIMIT:
                total += OMML_WEIGHT_MEDIUM
            else:
                total += OMML_WEIGHT_LONG
    return total

def get_option_length_class(options_dict):
    if not options_dict: return "fourline"
    lens = [_option_effective_length(p, l) for l, (p, _) in options_dict.items() if p]
    if not lens: return "fourline"
    if all(l <= SHORT_OPTION_CHAR_LIMIT for l in lens): return "oneline"
    if all(l <= MEDIUM_OPTION_CHAR_LIMIT for l in lens): return "twoline"
    return "fourline"

def parse_serial_and_question(line):
    m = re.match(r'^\s*[\(]?([০-৯0-9]+)[\.\)\।]?\s*(.*)', line)
    return m.groups() if m else (None, line)

def parse_option(line):
    m = re.match(r'^\s*[\(\[]?\s*([কখগঘa-dA-D])[\.\)\]\।]?\s*(.*)', line)
    return (m.group(1), m.group(2).strip()) if m else (None, line)

def parse_answer(line):
    m = re.match(r'^উত্তর[:：ঃ]?\s*[\(\[]?\s*([কখগঘa-dA-D])[\.\)\]\।]?\s*(.*)', line, re.IGNORECASE)
    if m:
        return (m.group(1), m.group(2).lstrip(').।. ').strip())
    m2 = re.match(r'^উত্তর[:：ঃ]?\s*(.*)', line, re.IGNORECASE)
    return (None, m2.group(1).lstrip(').।. ').strip()) if m2 else (None, None)

def parse_explanation(line):
    m = re.match(r'^(ব্যাখ্যা[:：ঃ]?)\s*(.*)', line)
    return (m.group(1), m.group(2)) if m else (None, None)


def extract_mcqs(paragraphs):
    """Extracts all structured MCQ data from a list of paragraphs."""
    mcqs = []
    para_to_mcq = {}
    cur = None
    state = "find_question"

    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        line = para.text.strip()
        
        if not line and state != "in_options":
            i += 1
            continue

        if state == "find_question":
            serial, _ = parse_serial_and_question(line)
            if serial:
                cur = {
                    'serial': serial, 'question_meta': [(para, None)], 'reference': '', 'qtype': '',
                    'options_meta': {}, 'answer_label': None, 'answer_text': None,
                    'explanation_meta': None, 'all_paras': {i}
                }
                para_to_mcq[i] = len(mcqs)
                state = "in_question"
            i += 1
        
        elif state == "in_question":
            opt_label, _ = parse_option(line)
            if opt_label:
                state = "in_options"
                continue
            
            # This is part of the question
            mref = re.search(r'(\[.*?\])', line)
            if mref: cur['reference'] += clean_text(mref.group(1))
            cur['question_meta'].append((para, None))
            cur['all_paras'].add(i)
            para_to_mcq[i] = len(mcqs)
            i += 1

        elif state == "in_options":
            opt_label, opt_text = parse_option(line)
            ans_label, ans_text = parse_answer(line)
            exp_label, exp_text = parse_explanation(line)

            if opt_label:
                cur['options_meta'][opt_label] = (para, opt_text)
            elif ans_label is not None or ans_text is not None:
                cur['answer_label'] = ans_label
                cur['answer_text'] = ans_text
            elif exp_label:
                cur['explanation_meta'] = (para, exp_label)
            
            cur['all_paras'].add(i)
            para_to_mcq[i] = len(mcqs)
            
            if ans_label is not None or ans_text is not None or exp_label is not None:
                if not exp_label and (i + 1) < len(paragraphs):
                    next_para = paragraphs[i+1]
                    if parse_explanation(next_para.text.strip())[0]:
                        cur['explanation_meta'] = (next_para, parse_explanation(next_para.text.strip())[0])
                        cur['all_paras'].add(i + 1)
                mcqs.append(cur)
                state = "find_question"
                cur = None
            i += 1
            
    if cur: mcqs.append(cur)
    return mcqs, para_to_mcq

# --- Document Formatting and Generation ---

def format_mcq(mcq, doc, tab_stops_oneline, tab_stops_twoline):
    """Formats and writes a single MCQ to the output document."""
    
    # --- Question ---
    q_para = doc.add_paragraph()
    set_bangla_font(q_para.add_run(f"{mcq['serial']}. "), bold=True)
    
    # Question can span multiple paragraphs
    full_question_parts = []
    prompt_line_para = None
    roman_numeral_paras = []

    for i, (p, _) in enumerate(mcq['question_meta']):
        text = p.text.strip()
        if i == 0:
            # Main question line, strip the original serial
            full_question_parts.extend(split_text_and_omml(p, strip_label_prefix=mcq['serial']))
        else:
            if any(key in text for key in ["সঠিক", "ঠিক", "যথাযথ", " কোনটি", "কোনগুলো"]):
                prompt_line_para = p
            else: # It's a roman numeral statement
                roman_numeral_paras.append(p)
    
    render_parts_to_para(q_para, full_question_parts)
    
    # Add reference tags to the main question paragraph
    ref = mcq.get('reference', '').strip()
    if ref:
        q_para.add_run("  ")
        set_bangla_font(q_para.add_run(ref))

    # Render roman numeral items
    for i, p in enumerate(roman_numeral_paras):
        sub_p = doc.add_paragraph()
        # set_bangla_font(sub_p.add_run(f"\t{ROMAN_NUMERALS[i]} "), bold=False)
        set_bangla_font(sub_p.add_run(f"{ROMAN_NUMERALS[i]} "), bold=True)
        render_parts_to_para(sub_p, split_text_and_omml(p))

    # Render the prompt line (e.g., "Which is correct?")
    if prompt_line_para:
        render_parts_to_para(doc.add_paragraph(), split_text_and_omml(prompt_line_para))

    # --- Options ---
    # ... (rest of the function is the same, but will now work correctly
    #      because split_text_and_omml and render_parts_to_para are fixed)
    option_layout = get_option_length_class(mcq['options_meta'])
    label_order = list(mcq['options_meta'].keys())
    
    def render_option(p, label, bn_label):
        opt_tuple = mcq['options_meta'].get(label)
        if not opt_tuple: return
        set_bangla_font(p.add_run(f"{bn_label}. "), bold=True)
        parts = split_text_and_omml(opt_tuple[0], strip_label_prefix=label)
        render_parts_to_para(p, parts)
        
    if option_layout == "oneline":
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.clear_all()
        for tab in tab_stops_oneline: p.paragraph_format.tab_stops.add_tab_stop(tab)
        for i, label in enumerate(label_order[:4]):
            if i > 0: p.add_run('\t')
            render_option(p, label, OPTION_LABELS_BN[i])
            
    elif option_layout == "twoline":
        for i, row in enumerate([(0, 1), (2, 3)]):
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.clear_all()
            for tab in tab_stops_twoline: p.paragraph_format.tab_stops.add_tab_stop(tab)
            for j, opt_idx in enumerate(row):
                if opt_idx >= len(label_order): continue
                if j > 0: p.add_run('\t')
                render_option(p, label_order[opt_idx], OPTION_LABELS_BN[opt_idx])
    else: # "fourline"
        for i, label in enumerate(label_order[:4]):
            p = doc.add_paragraph()
            render_option(p, label, OPTION_LABELS_BN[i])

    # --- Answer section ---
    ans_label = mcq.get('answer_label')
    ans_text = (mcq.get('answer_text') or "").strip()

    if ans_label or ans_text:
        p = doc.add_paragraph()
        set_bangla_font(p.add_run("উত্তর: "), bold=True, color=ANSWER_COLOR)
        if ans_label:
            # "খ" found as label, print label + answer
            set_bangla_font(p.add_run(f"{ans_label}. "))        # set_bangla_font(p.add_run(f"{ans_label}. "), color=ANSWER_COLOR)
            ans_tuple = mcq['options_meta'].get(ans_label)
            if ans_tuple:
                parts = split_text_and_omml(ans_tuple[0], strip_label_prefix=ans_label)
                render_parts_to_para(p, parts)                  # render_parts_to_para(p, parts, color=ANSWER_COLOR)
            elif ans_text:
                set_bangla_font(p.add_run(ans_text))            # set_bangla_font(p.add_run(ans_text), color=ANSWER_COLOR)
        elif ans_text:
            # Try to match the answer text to one of the options (normalized)
            found_label = None
            norm_ans = ans_text.replace(" ", "").replace("।", "").strip()
            for label, (para, opt_text) in mcq['options_meta'].items():
                norm_opt = opt_text.replace(" ", "").replace("।", "").strip()
                if norm_ans == norm_opt:
                    found_label = label
                    break
            if found_label:
                set_bangla_font(p.add_run(f"{found_label}. "))  # set_bangla_font(p.add_run(f"{found_label}. "), color=ANSWER_COLOR)
                set_bangla_font(p.add_run(ans_text))            # set_bangla_font(p.add_run(ans_text), color=ANSWER_COLOR)
            else:
                set_bangla_font(p.add_run(ans_text))            # set_bangla_font(p.add_run(ans_text), color=ANSWER_COLOR)
        p.paragraph_format.space_after = Pt(8)

    # --- Explanation ---
    if mcq.get('explanation_meta'):
        para, label = mcq['explanation_meta']
        ep = doc.add_paragraph()
        set_bangla_font(ep.add_run("ব্যাখ্যা: "), bold=True)
        parts = split_text_and_omml(para, strip_label_prefix=label)
        render_parts_to_para(ep, parts)

    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)


def copy_para_with_omml(src_para, outdoc, bold=False):
    out_p = outdoc.add_paragraph()
    render_parts_to_para(out_p, split_text_and_omml(src_para), bold=bold)
    return out_p

def is_extra_heading(text):
    keywords = ["উদ্দীপক", "নিচের", "তথ্যের ভিত্তিতে", "তথ্যের আলোকে", "প্রশ্নের উত্তর দাও"]
    return any(text.strip().startswith(kw) for kw in keywords)

def convert_file(src_file, out_file):
    try:
        doc = Document(src_file)
    except Exception as e:
        messagebox.showerror("File Error", f"Could not open source file.\n\n{e}")
        return False
        
    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '210')

    tab_stops_oneline = [Inches(0.8), Inches(1.6), Inches(2.4)]
    tab_stops_twoline = [Inches(1.6)]

    mcqs, para_to_mcq = extract_mcqs(doc.paragraphs)
    processed_mcq_indices = set()
    
    i = 0
    while i < len(doc.paragraphs):
        if i in para_to_mcq:
            mcq_idx = para_to_mcq[i]
            if mcq_idx not in processed_mcq_indices:
                format_mcq(mcqs[mcq_idx], outdoc, tab_stops_oneline, tab_stops_twoline)
                processed_mcq_indices.add(mcq_idx)
            i = max(mcqs[mcq_idx]['all_paras']) + 1
        else:
            para = doc.paragraphs[i]
            if para.text.strip():
                copy_para_with_omml(para, outdoc, bold=is_extra_heading(para.text))
            i += 1
            
    try:
        outdoc.save(out_file)
    except Exception as e:
        messagebox.showerror("Save Error", f"Could not save output file.\n\n{e}")
        return False

    return True

# --- Tkinter GUI (Unchanged) ---
def main_gui():
    def select_input():
        filename = filedialog.askopenfilename(title="Select DOCX file", filetypes=[("Word Files", "*.docx")])
        if filename:
            input_entry.delete(0, tk.END)
            input_entry.insert(0, filename)

    def select_output():
        filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")], title="Save Reformatted DOCX as")
        if filename:
            output_entry.delete(0, tk.END)
            output_entry.insert(0, filename)

    def run_conversion():
        src_file = input_entry.get().strip()
        out_file = output_entry.get().strip()
        if not src_file or not out_file:
            messagebox.showwarning("Input Needed", "Please select both input and output files!")
            return
        
        root.config(cursor="wait")
        root.update()
        try:
            ok = convert_file(src_file, out_file)
        finally:
            root.config(cursor="")
        
        if ok:
            messagebox.showinfo("Done!", f"Successfully converted and saved to:\n{out_file}")

    root = tk.Tk()
    root.title("MCQ Sheet Formatter (.docx)")

    tk.Label(root, text="Input DOCX:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
    input_entry = tk.Entry(root, width=50)
    input_entry.grid(row=0, column=1, padx=5)
    tk.Button(root, text="Browse...", command=select_input).grid(row=0, column=2, padx=5)

    tk.Label(root, text="Output DOCX:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
    output_entry = tk.Entry(root, width=50)
    output_entry.grid(row=1, column=1, padx=5)
    tk.Button(root, text="Browse...", command=select_output).grid(row=1, column=2, padx=5)

    tk.Button(root, text="Convert File", command=run_conversion, bg="#13825c", fg="white", width=16, height=2).grid(row=2, column=1, pady=15)

    root.mainloop()

if __name__ == "__main__":
    main_gui()