from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches
import re

def main():
    src = "AP_-_MCQ_Sheet_-_Class_6_-_Chapter_1.1^J_1.2^J_1.3^J_1.4^J_1.5^J_1.6_-_স্বাভাবিক_সংখ্যা_ও_ভগ্নাংশ.docx"
    out = "Reformatted_MCQ_Sheet.docx"
    doc = Document(src)

    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.6)
    section.footer_distance = Inches(0.2)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # Set to 2 columns

    # Parse MCQ blocks
    mcqs = parse_mcqs_from_paragraphs(doc.paragraphs)
    for mcq in mcqs:
        format_mcq(mcq, outdoc)

    outdoc.save(out)
    print("Done! Output:", out)

def parse_mcqs_from_paragraphs(paragraphs):
    # Simple state machine for MCQ detection
    mcqs = []
    cur = {}
    state = 0  # 0=question, 1=options, 2=answer
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Question: starts with digit dot
        if re.match(r'^\d+\.', text):
            if cur:
                mcqs.append(cur)
            cur = {
                'serial': text.split('.')[0],
                'question_para': para,
                'options': [],
                'answer_para': None
            }
            state = 1
        # Options: starts with Bangla "ক." "খ." "গ." "ঘ."
        elif re.match(r'^[কখগঘ]\.', text):
            cur['options'].append(para)
        # Answer: starts with "উত্তর:"
        elif text.startswith('উত্তর:'):
            cur['answer_para'] = para
    if cur:
        mcqs.append(cur)
    return mcqs

def copy_paragraph_with_equations_and_style(src_para, out_para, font_name="Tiro Bangla", font_size=11, color_hex=None, bold=False):
    from lxml import etree
    para_xml = etree.tostring(src_para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            texts = node.findall('.//w:t', namespaces=ns)
            for t in texts:
                run = out_para.add_run(t.text)
                run.font.name = font_name
                run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = Pt(font_size)
                if color_hex:
                    run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))
                run.font.bold = bold
        elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
            omml_xml = etree.tostring(node, encoding='unicode')
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_xml +
                '</w:r>'
            )
            out_para._p.append(omml_run)

def format_mcq(mcq, doc):
    # Question
    q_para = doc.add_paragraph()
    copy_paragraph_with_equations_and_style(mcq['question_para'], q_para)

    # Options: always two per line, preserving OMML equations
    opts = mcq['options']
    def get_opt_text_para(para):
        temp = doc.add_paragraph()
        copy_paragraph_with_equations_and_style(para, temp)
        return temp

    if len(opts) == 4:
        # First line: ক, খ
        opt_line1 = doc.add_paragraph()
        copy_paragraph_with_equations_and_style(opts[0], opt_line1)
        opt_line1.add_run('\t\t')
        copy_paragraph_with_equations_and_style(opts[1], opt_line1)
        # Second line: গ, ঘ
        opt_line2 = doc.add_paragraph()
        copy_paragraph_with_equations_and_style(opts[2], opt_line2)
        opt_line2.add_run('\t\t')
        copy_paragraph_with_equations_and_style(opts[3], opt_line2)
    else:
        # fallback: print each option in a separate line
        for para in opts:
            op = doc.add_paragraph()
            copy_paragraph_with_equations_and_style(para, op)

    # Answer
    ans_para = doc.add_paragraph()
    if mcq['answer_para']:
        # Color only the "উত্তর:" part
        text = mcq['answer_para'].text
        if text.startswith("উত্তর:"):
            run = ans_para.add_run("উত্তর:")
            run.font.name = "Tiro Bangla"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x08, 0x85, 0x65)
            rest = text[len("উত্তর:"):]
            run2 = ans_para.add_run(rest)
            run2.font.name = "Tiro Bangla"
            run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
            run2.font.size = Pt(11)
        else:
            copy_paragraph_with_equations_and_style(mcq['answer_para'], ans_para, color_hex="#088565", bold=True)
    # doc.add_paragraph('')  # Blank line for separation

if __name__ == "__main__":
    main()
