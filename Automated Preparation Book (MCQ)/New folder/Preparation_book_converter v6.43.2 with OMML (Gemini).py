import re
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches

# =============================================================================
# CONFIGURATION
# =============================================================================
# Centralize all settings here for easy modification.

CONFIG = {
    "files": {
        "source": "AP_-_MCQ_Sheet_-_Class_6_-_Chapter_1.1^J_1.2^J_1.3^J_1.4^J_1.5^J_1.6_-_স্বাভাবিক_সংখ্যা_ও_ভগ্নাংশ.docx",
        "output": "Reformatted_MCQ_Sheet_Professional.docx",
    },
    "font": {
        "name": "Tiro Bangla",
        "size_pt": 11,
    },
    "layout": {
        "columns": 2,
        "page_width_in": 8.5,
        "page_height_in": 10.65,
        "margin_top_in": 0.5,
        "margin_bottom_in": 0.3,
        "margin_left_in": 0.8,
        "margin_right_in": 0.6,
    },
    "styles": {
        "answer_color_rgb": RGBColor(0x08, 0x85, 0x65), # A shade of green
    },
    "options": {
        "labels": ['ক', 'খ', 'গ', 'ঘ'],
        "long_option_char_limit": 22, # Options longer than this will force a single-column layout
    },
    "regex": {
        # Matches the question number, e.g., "1."
        "question": re.compile(r'^\s*(\d+)\.\s*(.*)'),
        # Matches a reference, e.g., "[Some Board '23]"
        "reference": re.compile(r'\s*\[(.+)\]\s*'),
        # Matches an option, e.g., "ক. Some text"
        "option": re.compile(r'^\s*(ক|খ|গ|ঘ)\.\s*(.*)'),
        # Matches the answer line, e.g., "উত্তর: ক"
        "answer": re.compile(r'^\s*উত্তর[:：]\s*(\w)\.?\s*(.*)'),
    }
}


# =============================================================================
# LOW-LEVEL UTILITY FUNCTIONS
# =============================================================================

def _patch_omml_font_size(omml_xml, size_pt):
    """
    Modifies the font size within an OMML (Office Math Markup Language) string.

    Args:
        omml_xml (str): The XML string of the math equation.
        size_pt (int): The target font size in points.

    Returns:
        str: The modified OMML XML string.
    """
    try:
        size_val = str(int(size_pt * 2))  # Font size in half-points
        tree = etree.fromstring(omml_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Find and update all font size tags
        for sz_tag in tree.xpath('.//w:sz|.//w:szCs', namespaces=ns):
            sz_tag.attrib[qn('w:val')] = size_val
            
        return etree.tostring(tree, encoding='unicode')
    except etree.XMLSyntaxError:
        # Return original XML if parsing fails
        return omml_xml


def _split_paragraph_content(para):
    """
    Splits a paragraph into a list of text strings and OMML XML strings.

    Args:
        para (docx.paragraph.Paragraph): The source paragraph.

    Returns:
        list: A list of tuples, where each tuple is ('text', content) or ('omml', content).
    """
    para_xml = etree.tostring(para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    content_parts = []
    
    # Namespace map for parsing WordprocessingML
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    }

    for node in tree.iterchildren():
        if node.tag.endswith('r'):  # A 'run' contains text
            text = "".join(t.text or "" for t in node.findall('.//w:t', namespaces=ns))
            if text:
                content_parts.append(('text', text))
        elif node.tag.endswith(('oMath', 'oMathPara')):  # An OMML math object
            omml_xml = etree.tostring(node, encoding='unicode')
            content_parts.append(('omml', omml_xml))
            
    return content_parts


# =============================================================================
# DOCUMENT WRITING HELPERS
# =============================================================================

def _add_styled_run(para, text, bold=False, color=None, font_name=CONFIG['font']['name'], size_pt=CONFIG['font']['size_pt']):
    """
    Adds a styled run of text to a paragraph.

    Args:
        para (docx.paragraph.Paragraph): The paragraph to add the run to.
        text (str): The text content.
        bold (bool): Whether the text should be bold.
        color (RGBColor): The color of the text.
        font_name (str): The name of the font.
        size_pt (int): The font size in points.
    """
    run = para.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_omml_to_para(para, omml_xml, size_pt=CONFIG['font']['size_pt']):
    """
    Adds an OMML math equation to a paragraph.

    Args:
        para (docx.paragraph.Paragraph): The paragraph to add the equation to.
        omml_xml (str): The OMML XML string of the equation.
        size_pt (int): The font size for the equation.
    """
    patched_omml = _patch_omml_font_size(omml_xml, size_pt)
    omml_run = parse_xml(
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{patched_omml}</w:r>'
    )
    para._p.append(omml_run)


# =============================================================================
# CORE LOGIC: PARSING & FORMATTING
# =============================================================================

def extract_mcqs(paragraphs):
    """
    Extracts all MCQs from a list of paragraphs from a docx file.

    Args:
        paragraphs (list): A list of docx.paragraph.Paragraph objects.

    Returns:
        list: A list of dictionaries, where each dictionary represents one MCQ.
    """
    mcqs = []
    current_mcq = {}
    state = "find_question" # States: find_question, find_options, find_answer

    q_re = CONFIG['regex']['question']
    opt_re = CONFIG['regex']['option']
    ans_re = CONFIG['regex']['answer']
    ref_re = CONFIG['regex']['reference']

    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue

        try:
            # --- State: Looking for the start of a new question ---
            if state == "find_question":
                match = q_re.match(text)
                if match:
                    current_mcq = {
                        'serial': match.group(1),
                        'question_meta': [(match.group(2), para)],
                        'reference': '',
                        'options_meta': {},
                        'answer_label': '',
                        'answer_text': ''
                    }
                    state = "find_options"
                continue

            # --- State: Looking for options or more question lines ---
            if state == "find_options":
                opt_match = opt_re.match(text)
                ref_match = ref_re.search(text)

                if opt_match:
                    # Found the first option, switch state and process it
                    state = "find_answer"
                    current_mcq['options_meta'][opt_match.group(1)] = (opt_match.group(2), para)
                elif ref_match:
                    # Found a reference, add it to the question
                    current_mcq['reference'] = ref_match.group(0).strip()
                    # Remove the reference from the question line itself
                    q_text, q_para = current_mcq['question_meta'][-1]
                    current_mcq['question_meta'][-1] = (q_re.sub('', q_text).strip(), q_para)
                else:
                    # This is another line of the question (e.g., a list item)
                    current_mcq['question_meta'].append((text, para))
                continue

            # --- State: Looking for more options or the final answer ---
            if state == "find_answer":
                ans_match = ans_re.match(text)
                opt_match = opt_re.match(text)

                if ans_match:
                    # Found the answer, finalize the MCQ and reset state
                    current_mcq['answer_label'] = ans_match.group(1)
                    # Get the answer text from the options we've stored
                    ans_tuple = current_mcq['options_meta'].get(current_mcq['answer_label'])
                    if ans_tuple:
                        current_mcq['answer_text'] = ans_tuple[0]
                    
                    mcqs.append(current_mcq)
                    state = "find_question"
                elif opt_match:
                    # Found another option
                    current_mcq['options_meta'][opt_match.group(1)] = (opt_match.group(2), para)
                else:
                    # This is a continuation of the previous option
                    if current_mcq['options_meta']:
                        last_label = list(current_mcq['options_meta'])[-1]
                        prev_text, prev_para = current_mcq['options_meta'][last_label]
                        current_mcq['options_meta'][last_label] = (f"{prev_text} {text}", prev_para)
        
        except Exception as e:
            print(f"Error processing text near '{text[:50]}...'. Skipping. Error: {e}")
            # Reset state to find the next valid question
            state = "find_question"
            current_mcq = {}

    return mcqs


def format_mcq(mcq, doc):
    """
    Formats and writes a single MCQ to the output document.

    Args:
        mcq (dict): The MCQ data dictionary.
        doc (docx.Document): The output document object.
    """
    # --- 1. Write the Question and Reference ---
    q_para = doc.add_paragraph()
    _add_styled_run(q_para, f"{mcq['serial']}. ", bold=True)
    
    # Process all question lines (the first one is handled slightly differently)
    first_q_text, first_q_para = mcq['question_meta'][0]
    
    # Use the original paragraph to preserve mixed text/equations
    for content_type, content in _split_paragraph_content(first_q_para):
        if content_type == 'text':
            # Remove the serial number from the text before adding
            cleaned_text = CONFIG['regex']['question'].sub('', content).strip()
            # Also remove reference if present in the same line
            cleaned_text = CONFIG['regex']['reference'].sub('', cleaned_text).strip()
            if cleaned_text:
                _add_styled_run(q_para, cleaned_text + ' ')
        elif content_type == 'omml':
            _add_omml_to_para(q_para, content)
    
    # Add reference if it exists
    if mcq.get('reference'):
        _add_styled_run(q_para, f" {mcq['reference']}")

    # Add any subsequent question lines (e.g., roman numeral lists)
    for q_text, q_para in mcq['question_meta'][1:]:
        sub_para = doc.add_paragraph()
        for c_type, c_content in _split_paragraph_content(q_para):
            if c_type == 'text':
                _add_styled_run(sub_para, c_content + ' ')
            elif c_type == 'omml':
                _add_omml_to_para(sub_para, c_content)

    # --- 2. Write the Options ---
    # Decide layout: single column if any option is too long or contains only equations
    is_long = False
    only_equations = []
    
    for label in CONFIG['options']['labels']:
        opt_tuple = mcq['options_meta'].get(label)
        if opt_tuple:
            opt_text, opt_para = opt_tuple
            if len(opt_text) > CONFIG['options']['long_option_char_limit']:
                is_long = True
            
            content_parts = _split_paragraph_content(opt_para)
            text_content = "".join(part[1] for part in content_parts if part[0] == 'text')
            # Clean label from text
            text_content = CONFIG['regex']['option'].sub('', text_content).strip()
            if not text_content and any(p[0] == 'omml' for p in content_parts):
                only_equations.append(True)

    # If any option is long or more than half are just equations, use single-column layout
    if is_long or sum(only_equations) >= 2:
        # Single-column layout
        for label in CONFIG['options']['labels']:
            opt_tuple = mcq['options_meta'].get(label)
            if not opt_tuple: continue
            
            p = doc.add_paragraph()
            _add_styled_run(p, f"{label}. ", bold=True)
            for c_type, c_content in _split_paragraph_content(opt_tuple[1]):
                if c_type == 'text':
                    cleaned_text = CONFIG['regex']['option'].sub('', c_content).strip()
                    if cleaned_text:
                         _add_styled_run(p, cleaned_text + ' ')
                elif c_type == 'omml':
                    _add_omml_to_para(p, c_content)

    else:
        # Two-column layout
        option_pairs = [
            (CONFIG['options']['labels'][0], CONFIG['options']['labels'][1]),
            (CONFIG['options']['labels'][2], CONFIG['options']['labels'][3])
        ]
        for p_idx, pair in enumerate(option_pairs):
            p = doc.add_paragraph()
            for label in pair:
                opt_tuple = mcq['options_meta'].get(label)
                if not opt_tuple: continue

                _add_styled_run(p, f"{label}. ", bold=True)
                for c_type, c_content in _split_paragraph_content(opt_tuple[1]):
                    if c_type == 'text':
                        cleaned_text = CONFIG['regex']['option'].sub('', c_content).strip()
                        if cleaned_text:
                            _add_styled_run(p, cleaned_text)
                    elif c_type == 'omml':
                        _add_omml_to_para(p, c_content)
                
                # Add tabs to separate options
                p.add_run('\t\t')

    # --- 3. Write the Answer ---
    ans_para = doc.add_paragraph()
    ans_label = mcq.get('answer_label', '')
    if ans_label:
        _add_styled_run(ans_para, "উত্তর: ", bold=True, color=CONFIG['styles']['answer_color_rgb'])
        _add_styled_run(ans_para, f"{ans_label}. ", bold=True, color=CONFIG['styles']['answer_color_rgb'])
        
        # Add the full answer content (text and/or equation)
        ans_tuple = mcq['options_meta'].get(ans_label)
        if ans_tuple:
            for c_type, c_content in _split_paragraph_content(ans_tuple[1]):
                if c_type == 'text':
                    cleaned_text = CONFIG['regex']['option'].sub('', c_content).strip()
                    if cleaned_text:
                         _add_styled_run(ans_para, cleaned_text + ' ')
                elif c_type == 'omml':
                    _add_omml_to_para(ans_para, c_content)
    
    # Add spacing after the entire MCQ block
    ans_para.paragraph_format.space_after = Pt(8)


def main():
    """
    Main function to run the entire document conversion process.
    """
    print(f"Starting conversion of '{CONFIG['files']['source']}'...")
    
    try:
        source_doc = Document(CONFIG['files']['source'])
    except Exception as e:
        print(f"FATAL ERROR: Could not open source file. Error: {e}")
        return

    # Create and configure the output document
    output_doc = Document()
    section = output_doc.sections[0]
    
    # Set page layout from CONFIG
    lo = CONFIG['layout']
    section.page_width = Inches(lo['page_width_in'])
    section.page_height = Inches(lo['page_height_in'])
    section.top_margin = Inches(lo['margin_top_in'])
    section.bottom_margin = Inches(lo['margin_bottom_in'])
    section.left_margin = Inches(lo['margin_left_in'])
    section.right_margin = Inches(lo['margin_right_in'])
    
    # Set to two columns
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(lo['columns']))

    # Extract and format MCQs
    mcqs = extract_mcqs(source_doc.paragraphs)
    print(f"Found {len(mcqs)} MCQs to process.")
    
    for mcq in mcqs:
        format_mcq(mcq, output_doc)

    # Save the final document
    try:
        output_doc.save(CONFIG['files']['output'])
        print(f"Success! Formatted document saved as '{CONFIG['files']['output']}'")
    except Exception as e:
        print(f"FATAL ERROR: Could not save output file. Error: {e}")


if __name__ == "__main__":
    main()