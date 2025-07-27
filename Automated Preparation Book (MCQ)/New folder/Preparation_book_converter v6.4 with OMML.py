from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches
import re

def patch_omml_font_size(omml_xml, size_pt=11):
    from lxml import etree
    size_val = str(int(size_pt * 2))
    tree = etree.fromstring(omml_xml)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for sz_tag in tree.xpath('.//w:sz', namespaces=ns):
        sz_tag.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = size_val
    for sz_tag in tree.xpath('.//w:szCs', namespaces=ns):
        sz_tag.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = size_val
    return etree.tostring(tree, encoding='unicode')

def set_tiro_bangla_font(para):
    for run in para.runs:
        run.font.name = "Tiro Bangla"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
        run.font.size = Pt(11)
    # Set line spacing and paragraph spacing
    p_format = para.paragraph_format
    p_format.line_spacing = 1.3
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

def copy_paragraph_with_equations_and_style(src_para, out_para, font_name="Tiro Bangla", font_size=11, color_hex=None, bold=False):
    from lxml import etree
    para_xml = etree.tostring(src_para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            texts = node.findall('.//w:t', namespaces=ns)
            for t in texts:
                run = out_para.add_run(t.text)
                run.font.name = font_name
                run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = Pt(font_size)
                if color_hex:
                    run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))
                run.font.bold = bold
        elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
            omml_xml = etree.tostring(node, encoding='unicode')
            omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=font_size)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_xml_patched +
                '</w:r>'
            )
            out_para._p.append(omml_run)
    p_format = out_para.paragraph_format
    p_format.line_spacing = 1.3
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    return out_para

def format_mcq(mcq, doc):
    # Serial + Question (first line)
    qline0, style0, src_para0 = mcq['question_lines_meta'][0]
    serial = mcq['serial']

    # Add serial only once and not in the copied question text
    q_para = doc.add_paragraph()
    run = q_para.add_run(f"{serial}. ")
    run.font.name = "Tiro Bangla"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run.font.size = Pt(11)

    # Copy the question text but REMOVE the serial part if present
    from lxml import etree
    para_xml = etree.tostring(src_para0._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
          'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            texts = node.findall('.//w:t', namespaces=ns)
            for t in texts:
                text = t.text
                # Remove serial from start if present (e.g., "১. ")
                if text and text.strip().startswith(f"{serial}."):
                    text = text.strip()[len(f"{serial}."):].lstrip()
                run_q = q_para.add_run(text)
                run_q.font.name = "Tiro Bangla"
                run_q.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
                run_q.font.size = Pt(11)
        elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
            omml_xml = etree.tostring(node, encoding='unicode')
            omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=11)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_xml_patched +
                '</w:r>'
            )
            q_para._p.append(omml_run)

    # List-style sublines (roman) or extra text
    list_index = 1
    roman = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']
    for qline, style, src_para in mcq['question_lines_meta'][1:]:
        if 'List' in style:
            prefix = roman[list_index-1] if list_index <= len(roman) else f"{list_index}."
            p = doc.add_paragraph(prefix + " ")
            copy_paragraph_with_equations_and_style(src_para, p)
            list_index += 1
        elif qline.strip():
            p = doc.add_paragraph()
            copy_paragraph_with_equations_and_style(src_para, p)

    # Reference (if any)
    if mcq.get('reference'):
        ref_para = doc.add_paragraph(mcq['reference'])
        set_tiro_bangla_font(ref_para)

    # Options: two per line, remove label if already present in original
    opts = []
    for k in ['ক', 'খ', 'গ', 'ঘ']:
        v = mcq['options'].get(k)
        if v:
            opt_text, src_para = v
            opts.append((k, opt_text, src_para))
        else:
            opts.append((k, "", None))
    # First line: ক, খ
    opt_line1 = doc.add_paragraph()
    for idx in [0, 1]:
        label, opt_text, src_para = opts[idx]
        run_label = opt_line1.add_run(f"{label}. ")
        run_label.font.name = "Tiro Bangla"
        run_label.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
        run_label.font.size = Pt(11)
        if src_para:
            para_xml = etree.tostring(src_para._element, encoding='unicode')
            tree = etree.fromstring(para_xml.encode('utf-8'))
            for node in tree.iterchildren():
                if node.tag.endswith('r'):
                    texts = node.findall('.//w:t', namespaces=ns)
                    for t in texts:
                        text = t.text
                        # Remove label (e.g., "ক. ") from start
                        if text and text.strip().startswith(f"{label}."):
                            text = text.strip()[len(f"{label}."):].lstrip()
                        run = opt_line1.add_run(text)
                        run.font.name = "Tiro Bangla"
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
                        run.font.size = Pt(11)
                elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
                    omml_xml = etree.tostring(node, encoding='unicode')
                    omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=11)
                    omml_run = parse_xml(
                        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                        omml_xml_patched +
                        '</w:r>'
                    )
                    opt_line1._p.append(omml_run)
        else:
            run = opt_line1.add_run(opt_text)
            run.font.name = "Tiro Bangla"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
            run.font.size = Pt(11)
        opt_line1.add_run("\t\t")
    # Second line: গ, ঘ
    opt_line2 = doc.add_paragraph()
    for idx in [2, 3]:
        label, opt_text, src_para = opts[idx]
        run_label = opt_line2.add_run(f"{label}. ")
        run_label.font.name = "Tiro Bangla"
        run_label.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
        run_label.font.size = Pt(11)
        if src_para:
            para_xml = etree.tostring(src_para._element, encoding='unicode')
            tree = etree.fromstring(para_xml.encode('utf-8'))
            for node in tree.iterchildren():
                if node.tag.endswith('r'):
                    texts = node.findall('.//w:t', namespaces=ns)
                    for t in texts:
                        text = t.text
                        if text and text.strip().startswith(f"{label}."):
                            text = text.strip()[len(f"{label}."):].lstrip()
                        run = opt_line2.add_run(text)
                        run.font.name = "Tiro Bangla"
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
                        run.font.size = Pt(11)
                elif node.tag.endswith('oMath') or node.tag.endswith('oMathPara'):
                    omml_xml = etree.tostring(node, encoding='unicode')
                    omml_xml_patched = patch_omml_font_size(omml_xml, size_pt=11)
                    omml_run = parse_xml(
                        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                        omml_xml_patched +
                        '</w:r>'
                    )
                    opt_line2._p.append(omml_run)
        else:
            run = opt_line2.add_run(opt_text)
            run.font.name = "Tiro Bangla"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), "Tiro Bangla")
            run.font.size = Pt(11)
        opt_line2.add_run("\t\t")

    # Answer
    ans = mcq.get('answer', '')
    ans_text = mcq.get('answer_text', '')
    ans_para = doc.add_paragraph()
    run = ans_para.add_run("উত্তর:")
    run.font.name = 'Tiro Bangla'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x08, 0x85, 0x65)
    run2 = ans_para.add_run(f" {ans}. {ans_text}")
    run2.font.name = 'Tiro Bangla'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run2.font.size = Pt(11)
    # doc.add_paragraph('')


def extract_mcqs(paragraphs):
    # This version expects question_lines_meta with OMML
    mcqs = []
    q_re = re.compile(r'^(\d+)\.\s*(.*)')
    ref_re = re.compile(r'^\[(.+)\]')
    opt_re = re.compile(r'^(ক|খ|গ|ঘ)\.\s*(.+)')
    ans_re = re.compile(r'^উত্তর[:：]\s*(\w)\.?\s*(.*)')
    state = 0
    cur = {}
    for para in paragraphs:
        text = para.text.strip()
        style = para.style.name if hasattr(para, "style") else ""
        if not text:
            continue

        # --- Split each paragraph by new line and process each line ---
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue


        if state == 0:
            m = q_re.match(text)
            if m:
                cur = {
                    'serial': m.group(1),
                    'question_lines_meta': [(m.group(2), style, para)],
                    'reference': '',
                    'options': {},
                    'answer': ''
                }
                state = 1
        elif state == 1:
            m = ref_re.match(text)
            if m:
                cur['reference'] = m.group(0)
                continue
            if opt_re.match(text):
                state = 2
            else:
                cur['question_lines_meta'].append((text, style, para))
                continue
        if state == 2:
            m = opt_re.match(text)
            if m:
                cur['options'][m.group(1)] = (m.group(2), para)
                continue
            elif ans_re.match(text):
                m = ans_re.match(text)
                cur['answer'] = m.group(1)
                cur['answer_text'] = cur['options'].get(cur['answer'], ('', ''))[0]
                mcqs.append(cur)
                state = 0
            else:
                if cur['options']:
                    last_opt = list(cur['options'])[-1]
                    val, oxmls = cur['options'][last_opt]
                    cur['options'][last_opt] = (val + ' ' + text, para)
    return mcqs

def main():
    src = "AP_-_MCQ_Sheet_-_Class_6_-_Chapter_1.1^J_1.2^J_1.3^J_1.4^J_1.5^J_1.6_-_স্বাভাবিক_সংখ্যা_ও_ভগ্নাংশ.docx"
    out = "Reformatted_MCQ_Sheet.docx"
    doc = Document(src)

    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.6)
    section.footer_distance = Inches(0.2)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # Set to 2 columns

    mcqs = extract_mcqs(doc.paragraphs)
    for mcq in mcqs:
        format_mcq(mcq, outdoc)

    outdoc.save(out)
    print("Done! Output:", out)

if __name__ == "__main__":
    main()
