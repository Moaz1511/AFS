import re
import tkinter as tk
from tkinter import filedialog, messagebox
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches

FONT_NAME = "Tiro Bangla"
FONT_SIZE = 11
ANSWER_COLOR = RGBColor(0x08, 0x85, 0x65)
OPTION_LABELS_BN = ['ক', 'খ', 'গ', 'ঘ']
OPTION_LABELS_EN = ['a', 'b', 'c', 'd', 'A', 'B', 'C', 'D']
SHORT_OPTION_CHAR_LIMIT = 5
MEDIUM_OPTION_CHAR_LIMIT = 15
LONG_OPTION_CHAR_LIMIT = 20
OMML_WEIGHT_SHORT = 5
OMML_WEIGHT_MEDIUM = 12
OMML_WEIGHT_LONG = 1000
ROMAN_NUMERALS = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']

def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'([(\[\{«“‘])\s+', r'\1', text)
    text = re.sub(r'\s+([)\]\}»”’.,?।;:])', r'\1', text)
    text = re.sub(r' {2,}', ' ', text)
    return text

def patch_omml_font_size(omml_xml, size_pt=FONT_SIZE):
    try:
        size_val = str(int(size_pt * 2))
        tree = etree.fromstring(omml_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for sz_tag in tree.xpath('.//w:sz|.//w:szCs', namespaces=ns):
            sz_tag.attrib[qn('w:val')] = size_val
        return etree.tostring(tree, encoding='unicode')
    except etree.XMLSyntaxError:
        return omml_xml

def set_bangla_font(run, font_size=FONT_SIZE, bold=False, color=None):
    run.font.name = FONT_NAME
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def split_text_and_omml(para, strip_label_prefix=None):
    para_xml = etree.tostring(para._element, encoding='unicode')
    tree = etree.fromstring(para_xml.encode('utf-8'))
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    parts = []
    first_text_run = True
    for node in tree.iterchildren():
        if node.tag.endswith('r'):
            text = "".join(t.text or "" for t in node.findall('.//w:t', namespaces=ns))
            if first_text_run and strip_label_prefix:
                # Remove ONLY the label and its attached punctuations/spaces (like "(ক)", "ক.", "ক)", "ক।", "ক ", etc)
                regex = re.compile(r'^\s*[\(\[]?%s[\.\)\]\।]?\s*' % re.escape(strip_label_prefix))
                text = regex.sub('', text, count=1)
                first_text_run = False
            if text:
                parts.append(("text", text))
        elif node.tag.endswith(('oMath', 'oMathPara')):
            omml_xml = etree.tostring(node, encoding='unicode')
            parts.append(("omml", omml_xml))
    return parts


def render_parts_to_para(para, parts, bold=False, color=None):
    for i, (ctype, cvalue) in enumerate(parts):
        if i > 0:
            prev_ctype, _ = parts[i-1]
            if prev_ctype != ctype:
                para.add_run(" ")
        if ctype == "text":
            run = para.add_run(clean_text(cvalue))
            set_bangla_font(run, bold=bold, color=color)
        elif ctype == "omml":
            omml_patched = patch_omml_font_size(cvalue)
            omml_run = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">' +
                omml_patched +
                '</w:r>'
            )
            para._p.append(omml_run)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)

def _omml_visible_length(omml_xml):
    try:
        tree = etree.fromstring(omml_xml)
        math_texts = tree.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}t')
        text = ''.join([t.text or '' for t in math_texts])
        return len(text)
    except Exception:
        return 12

def _option_effective_length(para, label):
    total = 0
    for ctype, cval in split_text_and_omml(para, strip_label_prefix=label):
        if ctype == "text":
            total += len(cval.strip())
        elif ctype == "omml":
            visible_len = _omml_visible_length(cval)
            if visible_len <= SHORT_OPTION_CHAR_LIMIT:
                total += OMML_WEIGHT_SHORT
            elif visible_len <= MEDIUM_OPTION_CHAR_LIMIT:
                total += OMML_WEIGHT_MEDIUM
            else:
                total += OMML_WEIGHT_LONG
    return total

def get_option_length_class(options_dict):
    lens = []
    for label in options_dict:
        opt_tuple = options_dict.get(label)
        if not opt_tuple:
            return "fourline"
        para, _ = opt_tuple
        lens.append(_option_effective_length(para, label))
    if all(l <= SHORT_OPTION_CHAR_LIMIT for l in lens):
        return "oneline"
    elif all(l <= MEDIUM_OPTION_CHAR_LIMIT for l in lens):
        return "twoline"
    else:
        return "fourline"

def parse_serial_and_question(line):
    # Accepts ১., ১।, 1., (১), (1), ১), 1)
    m = re.match(r'^\s*[\(]?([০-৯0-9]+)[\.\)\।]?\s*(.*)', line)
    return m.groups() if m else (None, line)

def parse_option(line):
    m = re.match(r'^\s*[\(\[]?\s*([কখগঘa-dA-D])[\.\)\]\।]?\s*(.*)', line)
    if m:
        option_text = m.group(2).lstrip(').।. ')
        return (m.group(1), option_text)
    return (None, line)

def parse_answer(line):
    m = re.match(r'^উত্তর[:：ঃ]?\s*[\(\[]?\s*([কখগঘa-dA-D])[\.\)\]\।]?\s*(.*)', line)
    if m:
        answer_text = m.group(2).lstrip(').।. ')
        return (m.group(1), answer_text)
    m2 = re.match(r'^উত্তর[:：ঃ]?\s*(.*)', line)
    if m2:
        return (None, m2.group(1).lstrip(').।. '))
    return (None, None)

def parse_explanation(line):
    m = re.match(r'^(ব্যাখ্যা[:：ঃ]?)\s*(.*)', line)
    if m:
        return (m.group(1), m.group(2))
    return (None, None)

def extract_mcqs(paragraphs):
    mcqs = []
    para_to_mcq = {}
    cur = None
    state = "find_question"
    for i, para in enumerate(paragraphs):
        lines = para.text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if state == "find_question":
                serial, rest = parse_serial_and_question(line)
                if serial:
                    cur = {
                        'serial': serial,
                        'question_meta': [(para, rest)],
                        'reference': '',
                        'qtype': '',
                        'options_meta': {},
                        'answer_label': '',
                        'answer_text': '',
                        'explanation': None,
                        'all_paras': [i]
                    }
                    state = "in_question"
                    para_to_mcq[i] = len(mcqs)
            elif state == "in_question":
                mref = re.search(r'(\[.*?\])', line)
                mqtype = re.search(r'\[[^]]*(মূলক)[^]]*\]', line)
                opt_label, opt_text = parse_option(line)
                if mref:
                    cur['reference'] += clean_text(mref.group(1))
                if mqtype:
                    cur['qtype'] += clean_text(mqtype.group(0))
                if opt_label:
                    cur['options_meta'][opt_label] = (para, opt_text)
                    state = "in_options"
                elif line:
                    cur['question_meta'].append((para, line))
                cur['all_paras'].append(i)
            elif state == "in_options":
                ans_label, ans_text = parse_answer(line)
                opt_label, opt_text = parse_option(line)
                exp_label, exp_text = parse_explanation(line)
                if opt_label:
                    cur['options_meta'][opt_label] = (para, opt_text)
                elif ans_label or ans_text:
                    cur['answer_label'] = ans_label
                    cur['answer_text'] = ans_text
                elif exp_label:
                    cur['explanation'] = (para, exp_text)
                elif cur and cur['options_meta']:
                    last_label = list(cur['options_meta'])[-1]
                    prev_para, prev_text = cur['options_meta'][last_label]
                    cur['options_meta'][last_label] = (prev_para, f"{prev_text} {line}")
                cur['all_paras'].append(i)
        if cur and (cur.get('answer_label') or cur.get('answer_text')):
            # lookahead for explanation paragraph
            if i + 1 < len(paragraphs):
                next_text = paragraphs[i + 1].text.strip()
                exp_label, exp_text = parse_explanation(next_text)
                if exp_label:
                    cur['explanation'] = (paragraphs[i + 1], exp_text)
                    cur['all_paras'].append(i + 1)
            mcqs.append(cur)
            for idx in cur['all_paras']:
                para_to_mcq[idx] = len(mcqs) - 1
            cur = None
            state = "find_question"
    return mcqs, para_to_mcq

def format_mcq(mcq, doc, tab_stops_oneline, tab_stops_twoline):
    q_para = doc.add_paragraph()
    set_bangla_font(q_para.add_run(f"{mcq['serial']}. "), bold=True)
    prompt_line = None
    for i, (para, text) in enumerate(mcq['question_meta']):
        if i == 0:
            render_parts_to_para(q_para, split_text_and_omml(para, strip_label_prefix=mcq['serial']))
        else:
            if any(key in text for key in ["সঠিক", "ঠিক", "যথাযথ", " কোনটি", "কোনগুলো"]):
                prompt_line = para
            else:
                sub_para = doc.add_paragraph()
                set_bangla_font(sub_para.add_run(f"{ROMAN_NUMERALS[i-1]} "), bold=True)
                render_parts_to_para(sub_para, split_text_and_omml(para))
    if prompt_line:
        render_parts_to_para(doc.add_paragraph(), split_text_and_omml(prompt_line))
    refs = []
    if mcq.get('reference'):
        refs.append(mcq['reference'])
    if mcq.get('qtype'):
        refs.append(mcq['qtype'])
    if refs:
        q_para.add_run(" ")
        set_bangla_font(q_para.add_run(" ".join(refs)))
    option_layout = get_option_length_class(mcq['options_meta'])
    def all_labels():
        return [k for k in mcq['options_meta'].keys()]
    label_order = all_labels()
    label_order = label_order + [l for l in OPTION_LABELS_BN+OPTION_LABELS_EN if l not in label_order]
    if option_layout == "oneline":
        option_para = doc.add_paragraph()
        option_para.paragraph_format.tab_stops.clear_all()
        for tab_pos in tab_stops_oneline:
            option_para.paragraph_format.tab_stops.add_tab_stop(tab_pos)
        col = 0
        for label in label_order[:4]:
            opt_tuple = mcq['options_meta'].get(label)
            if not opt_tuple:
                continue
            if col > 0:
                option_para.add_run('\t')
            run = option_para.add_run(f"{label}. ")
            set_bangla_font(run, bold=True)
            render_parts_to_para(option_para, split_text_and_omml(opt_tuple[0], strip_label_prefix=label))
            col += 1
    elif option_layout == "twoline":
        for row in [(0, 1), (2, 3)]:
            option_para = doc.add_paragraph()
            option_para.paragraph_format.tab_stops.clear_all()
            for tab_pos in tab_stops_twoline:
                option_para.paragraph_format.tab_stops.add_tab_stop(tab_pos)
            col = 0
            for idx in row:
                label = label_order[idx]
                opt_tuple = mcq['options_meta'].get(label)
                if not opt_tuple:
                    continue
                if col > 0:
                    option_para.add_run('\t')
                run = option_para.add_run(f"{label}. ")
                set_bangla_font(run, bold=True)
                render_parts_to_para(option_para, split_text_and_omml(opt_tuple[0], strip_label_prefix=label))
                col += 1
    else:
        for idx in range(4):
            label = label_order[idx]
            opt_tuple = mcq['options_meta'].get(label)
            if not opt_tuple:
                continue
            p = doc.add_paragraph()
            run = p.add_run(f"{label}. ")
            set_bangla_font(run, bold=True)
            render_parts_to_para(p, split_text_and_omml(opt_tuple[0], strip_label_prefix=label))
    ans_label = mcq.get('answer_label')
    ans_text = mcq.get('answer_text', '').strip()
    if ans_label or ans_text:
        p = doc.add_paragraph()
        set_bangla_font(p.add_run("উত্তর: "), bold=True, color=ANSWER_COLOR)
        if ans_label:
            set_bangla_font(p.add_run(f"{ans_label}. "), bold=False)
            ans_tuple = mcq['options_meta'].get(ans_label)
            if ans_tuple:
                render_parts_to_para(p, split_text_and_omml(ans_tuple[0], strip_label_prefix=ans_label))
            elif ans_text:
                p.add_run(ans_text)
        elif ans_text:
            p.add_run(ans_text)
    if mcq.get('explanation'):
        para, exp_text = mcq['explanation']
        ep = doc.add_paragraph()
        set_bangla_font(ep.add_run("ব্যাখ্যা: "), bold=True)
        render_parts_to_para(ep, split_text_and_omml(para, strip_label_prefix="ব্যাখ্যা"))
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

def copy_para_with_omml(src_para, outdoc, bold=False):
    out_p = outdoc.add_paragraph()
    parts = split_text_and_omml(src_para)
    render_parts_to_para(out_p, parts, bold=bold)
    out_p.paragraph_format.space_after = Pt(0)
    out_p.paragraph_format.space_before = Pt(0)
    return out_p

def is_extra_heading(text):
    keywords = [
        "উদ্দীপক", "নিচের", "তথ্যের ভিত্তিতে", "তথ্যের আলোকে",
        "প্রশ্নের উত্তর দাও", "উত্তর দাও", "লক্ষ কর", "নিচের উদ্দীপকটি পড়ো"
    ]
    t = text.strip()
    return any(t.startswith(kw) for kw in keywords) and len(t) > 6

def convert_file(src_file, out_file):
    try:
        doc = Document(src_file)
    except Exception as e:
        messagebox.showerror("File Error", f"Could not open source file.\n\n{e}")
        return False
    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    tab_stops_oneline = [Inches(0.8), Inches(1.6), Inches(2.4)]
    tab_stops_twoline = [Inches(1.6)]
    mcqs, para_to_mcq = extract_mcqs(doc.paragraphs)
    i = 0
    while i < len(doc.paragraphs):
        if i in para_to_mcq:
            mcq_idx = para_to_mcq[i]
            if mcqs[mcq_idx].get('first_written', False) is False:
                format_mcq(mcqs[mcq_idx], outdoc, tab_stops_oneline, tab_stops_twoline)
                mcqs[mcq_idx]['first_written'] = True
            i = max(mcqs[mcq_idx]['all_paras']) + 1
        else:
            para_text = doc.paragraphs[i].text.strip()
            if not para_text or para_text.startswith("উত্তর:") or para_text.startswith("ব্যাখ্যা"):
                i += 1
                continue
            if is_extra_heading(para_text):
                copy_para_with_omml(doc.paragraphs[i], outdoc, bold=True)
            else:
                copy_para_with_omml(doc.paragraphs[i], outdoc, bold=False)
            i += 1
    outdoc.save(out_file)
    return True

# --- Tkinter GUI below ---
def main_gui():
    def select_input():
        filename = filedialog.askopenfilename(
            title="Select DOCX file",
            filetypes=[("Word Files", "*.docx")]
        )
        if filename:
            input_entry.delete(0, tk.END)
            input_entry.insert(0, filename)

    def select_output():
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Files", "*.docx")],
            title="Save Reformatted DOCX as"
        )
        if filename:
            output_entry.delete(0, tk.END)
            output_entry.insert(0, filename)

    def run_conversion():
        src_file = input_entry.get().strip()
        out_file = output_entry.get().strip()
        if not src_file or not out_file:
            messagebox.showwarning("Input Needed", "Please select both input and output files!")
            return
        root.config(cursor="wait")
        root.update()
        try:
            ok = convert_file(src_file, out_file)
        finally:
            root.config(cursor="")
        if ok:
            messagebox.showinfo("Done!", f"Successfully converted:\n{out_file}")

    root = tk.Tk()
    root.title("MCQ Sheet Formatter (.docx)")

    tk.Label(root, text="Input DOCX:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
    input_entry = tk.Entry(root, width=45)
    input_entry.grid(row=0, column=1, padx=5)
    tk.Button(root, text="Browse", command=select_input).grid(row=0, column=2, padx=5)

    tk.Label(root, text="Output DOCX:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
    output_entry = tk.Entry(root, width=45)
    output_entry.grid(row=1, column=1, padx=5)
    tk.Button(root, text="Browse", command=select_output).grid(row=1, column=2, padx=5)

    tk.Button(root, text="Convert", command=run_conversion, bg="#13825c", fg="white", width=16).grid(row=2, column=1, pady=15)

    root.mainloop()

if __name__ == "__main__":
    main_gui()
