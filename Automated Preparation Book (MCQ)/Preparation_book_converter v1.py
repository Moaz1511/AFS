from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
import re

def set_tiro_bangla_font(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Tiro Bangla'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
        run.font.size = Pt(11)
    # Set line spacing and paragraph spacing
    p_format = paragraph.paragraph_format
    p_format.line_spacing = 1.3
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

def color_ans_text(paragraph, color_hex='#088565'):
    # Only color "উত্তর:" in the answer line
    for run in paragraph.runs:
        if run.text.startswith('উত্তর:'):
            run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))
            break

def extract_mcqs(paragraphs):
    mcqs = []
    q_re = re.compile(r'^(\d+)\.\s*(.*)')
    ref_re = re.compile(r'^\[(.+)\]')
    opt_re = re.compile(r'^(ক|খ|গ|ঘ)\.\s*(.+)')
    ans_re = re.compile(r'^উত্তর[:：]\s*(\w)\.?\s*(.*)')
    
    state = 0  # 0: looking for question, 1: collecting question, 2: collecting options, 3: answer
    cur = {}
    question_lines = []
    for para in paragraphs:
        text = para.text.strip()
        if not text: continue
        if state == 0:
            m = q_re.match(text)
            if m:
                cur = {
                    'serial': m.group(1),
                    'question': '',     # Will collect all lines until options
                    'reference': '',
                    'options': {},
                    'answer': ''
                }
                question_lines = [m.group(2)]
                state = 1
        elif state == 1:
            # Look for reference
            m = ref_re.match(text)
            if m:
                cur['reference'] = m.group(0)
                continue
            # Look for option start
            if opt_re.match(text):
                cur['question'] = '\n'.join(question_lines)
                state = 2
            else:
                # Still part of question (subpoints, info, etc.)
                question_lines.append(text)
                continue
        if state == 2:
            m = opt_re.match(text)
            if m:
                cur['options'][m.group(1)] = m.group(2)
            elif ans_re.match(text):
                m = ans_re.match(text)
                cur['answer'] = m.group(1)
                cur['answer_text'] = cur['options'].get(cur['answer'], '')
                mcqs.append(cur)
                state = 0
    return mcqs


def format_mcq(mcq, doc):
    # 1. Serial + Question
    q_para = doc.add_paragraph()
    q_para.add_run(f"{mcq['serial']}. ")
    # Add the rest of the question, preserving line breaks
    for idx, line in enumerate(mcq['question'].split('\n')):
        if idx == 0:
            q_para.add_run(line)
        else:
            q_para.add_run('\n' + line)
    set_tiro_bangla_font(q_para)
    # 2. Reference (if any)
    if mcq['reference']:
        ref_para = doc.add_paragraph(f"{mcq['reference']}")
        set_tiro_bangla_font(ref_para)
    # 3. Options (dynamic layout)
    opts = [mcq['options'].get(k, '') for k in ['ক','খ','গ','ঘ']]
    max_opt_len = max([len(opt) for opt in opts if opt], default=0)
    # Decide if any option is "long"
    LONG_LIMIT = 14  # tune as needed
    if max_opt_len > LONG_LIMIT:
        # One per line
        for idx, opt in enumerate(['ক','খ','গ','ঘ']):
            p = doc.add_paragraph(f"{opt}. {mcq['options'].get(opt,'')}")
            set_tiro_bangla_font(p)
    else:
        # Two per line
        opt1 = f"ক. {mcq['options'].get('ক','')}"
        opt2 = f"খ. {mcq['options'].get('খ','')}"
        opt3 = f"গ. {mcq['options'].get('গ','')}"
        opt4 = f"ঘ. {mcq['options'].get('ঘ','')}"
        p1 = doc.add_paragraph(f"{opt1}\t\t{opt2}")
        set_tiro_bangla_font(p1)
        p2 = doc.add_paragraph(f"{opt3}\t\t{opt4}")
        set_tiro_bangla_font(p2)
    # 4. Answer (with option text, colorized)
    ans = mcq.get('answer', '')
    ans_text = mcq.get('answer_text', '')
    ans_para = doc.add_paragraph()
    run = ans_para.add_run(f"উত্তর:")  # color only this part
    run.font.name = 'Tiro Bangla'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x08, 0x85, 0x65)  # #088565
    run2 = ans_para.add_run(f" {ans}. {ans_text}")
    run2.font.name = 'Tiro Bangla'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run2.font.size = Pt(11)
    # 5. Blank line for separation
    # doc.add_paragraph('')

def main():
    src = "আলাউদ্দিনের চেরাগ MCQ.docx"
    out = "Reformatted_MCQ_Sheet.docx"
    doc = Document(src)
    mcqs = extract_mcqs(doc.paragraphs)

    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.6)
    section.footer_distance = Inches(0.2)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # Set to 2 columns

    for mcq in mcqs:
        format_mcq(mcq, outdoc)

    outdoc.save(out)
    print("Done! Output:", out)

if __name__ == "__main__":
    main()
