from docx import Document
from lxml import etree
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
import re

def set_tiro_bangla_font(paragraph):
    for run in paragraph.runs:
        run.font.name = 'Tiro Bangla'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
        run.font.size = Pt(11)
    # Set line spacing and paragraph spacing
    p_format = paragraph.paragraph_format
    p_format.line_spacing = 1.3
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

def color_ans_text(paragraph, color_hex='#088565'):
    # Only color "উত্তর:" in the answer line
    for run in paragraph.runs:
        if run.text.startswith('উত্তর:'):
            run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))
            break

def get_omml_xml(para):
    # Returns list of OMML XML strings in this paragraph
    omml_blocks = para._element.xpath('.//m:oMath | .//m:oMathPara')
    return [etree.tostring(omml, encoding="unicode") for omml in omml_blocks]

def extract_mcqs(paragraphs):
    mcqs = []
    q_re = re.compile(r'^(\d+)\.\s*(.*)')
    ref_re = re.compile(r'^\[(.+)\]')
    opt_re = re.compile(r'^(ক|খ|গ|ঘ)\.\s*(.+)')
    ans_re = re.compile(r'^উত্তর[:：]\s*(\w)\.?\s*(.*)')
    
    state = 0
    cur = {}
    for para in paragraphs:
        text = para.text.strip()
        style = para.style.name if hasattr(para, "style") else ""
        ommls = get_omml_xml(para)
        omml_info = ""
        # If equation present, append as placeholder in text for now
        if ommls:
            # For production: store ALL OMML as placeholder (convert to LaTeX/MathML later)
            omml_info = "[OMML_EQ] " + " ".join(ommls)
        if not text and ommls:
            text = omml_info
        if not text:
            continue
        if state == 0:
            m = q_re.match(text)
            if m:
                cur = {
                    'serial': m.group(1),
                    'question_lines_meta': [(m.group(2), style, ommls if ommls else None)],
                    'reference': '',
                    'options': {},
                    'answer': '',
                    'omml_equations': []
                }
                # Save any OMML found in first line
                if ommls:
                    cur['omml_equations'].extend(ommls)
                state = 1
        elif state == 1:
            m = ref_re.match(text)
            if m:
                cur['reference'] = m.group(0)
                continue
            if opt_re.match(text):
                state = 2
            else:
                cur['question_lines_meta'].append((text, style, ommls if ommls else None))
                if ommls:
                    cur['omml_equations'].extend(ommls)
                continue
        if state == 2:
            m = opt_re.match(text)
            if m:
                # Option may also have OMML
                option_text = m.group(2)
                # If option text is empty and OMML present, just save OMML
                if not option_text and ommls:
                    option_text = omml_info
                cur['options'][m.group(1)] = option_text
                if ommls:
                    cur['omml_equations'].extend(ommls)
                continue
            elif ans_re.match(text):
                m = ans_re.match(text)
                cur['answer'] = m.group(1)
                cur['answer_text'] = cur['options'].get(cur['answer'], '')
                mcqs.append(cur)
                state = 0
            else:
                # For option continuation lines, also check for OMML
                if cur['options']:
                    last_opt = list(cur['options'])[-1]
                    to_add = text
                    if ommls:
                        to_add += " [OMML_EQ] " + " ".join(ommls)
                        cur['omml_equations'].extend(ommls)
                    cur['options'][last_opt] += ' ' + to_add
    return mcqs


def format_mcq(mcq, doc):
    if not mcq['question_lines_meta']:
        print(f"Warning: MCQ {mcq.get('serial', '?')} has no question lines. Skipping.")
        return

    # 1. Serial + Question
    qline0, style0, ommls0 = mcq['question_lines_meta'][0]
    q_para = doc.add_paragraph(f"{mcq['serial']}. {qline0}")
    set_tiro_bangla_font(q_para)
    # Detect and number manually
    list_index = 1
    roman = ['i.', 'ii.', 'iii.', 'iv.', 'v.', 'vi.', 'vii.', 'viii.', 'ix.', 'x.']
    for qline, style, ommls in mcq['question_lines_meta'][1:]:
        if 'List' in style:
            prefix = roman[list_index-1] if list_index <= len(roman) else f"{list_index}."
            txt = f"{prefix} {qline}"
            if ommls:  # If there's OMML equation XML, show marker (or process later)
                txt += f"\n[OMML equation detected]"
            p = doc.add_paragraph(f"{prefix} {qline}")
            set_tiro_bangla_font(p)
            list_index += 1
        elif qline.strip():
            txt = qline
            if ommls:
                txt += f"\n[OMML equation detected]"
            p = doc.add_paragraph(qline)
            set_tiro_bangla_font(p)
    # 2. Reference (if any)
    if mcq.get('reference'):
        ref_para = doc.add_paragraph(f"{mcq['reference']}")
        set_tiro_bangla_font(ref_para)
    # 3. Options (dynamic layout)
    opts = [mcq['options'].get(k, '') for k in ['ক','খ','গ','ঘ']]
    max_opt_len = max([len(opt) for opt in opts if opt], default=0)
    LONG_LIMIT = 14
    if max_opt_len > LONG_LIMIT:
        for idx, opt in enumerate(['ক','খ','গ','ঘ']):
            p = doc.add_paragraph(f"{opt}. {mcq['options'].get(opt,'')}")
            set_tiro_bangla_font(p)
    else:
        opt1 = f"ক. {mcq['options'].get('ক','')}"
        opt2 = f"খ. {mcq['options'].get('খ','')}"
        opt3 = f"গ. {mcq['options'].get('গ','')}"
        opt4 = f"ঘ. {mcq['options'].get('ঘ','')}"
        p1 = doc.add_paragraph(f"{opt1}\t\t{opt2}")
        set_tiro_bangla_font(p1)
        p2 = doc.add_paragraph(f"{opt3}\t\t{opt4}")
        set_tiro_bangla_font(p2)
    # 4. Answer (with option text, colorized)
    ans = mcq.get('answer', '')
    ans_text = mcq.get('answer_text', '')
    ans_para = doc.add_paragraph()
    run = ans_para.add_run(f"উত্তর:")  # color only this part
    run.font.name = 'Tiro Bangla'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x08, 0x85, 0x65)  # #088565
    run2 = ans_para.add_run(f" {ans}. {ans_text}")
    run2.font.name = 'Tiro Bangla'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Tiro Bangla')
    run2.font.size = Pt(11)
    # 5. Blank line for separation
    # doc.add_paragraph('')

def main():
    src = "AP_-_MCQ_Sheet_-_Class_6_-_Chapter_1.1^J_1.2^J_1.3^J_1.4^J_1.5^J_1.6_-_স্বাভাবিক_সংখ্যা_ও_ভগ্নাংশ.docx"
    out = "Reformatted_MCQ_Sheet.docx"
    doc = Document(src)
    mcqs = extract_mcqs(doc.paragraphs)

    outdoc = Document()
    section = outdoc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(10.65)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.6)
    section.footer_distance = Inches(0.2)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')  # Set to 2 columns

    for mcq in mcqs:
        format_mcq(mcq, outdoc)

    outdoc.save(out)
    print("Done! Output:", out)

if __name__ == "__main__":
    main()
